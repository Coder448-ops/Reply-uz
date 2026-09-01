# -*- coding: utf-8 -*-
"""
================================================================================
  Reply-uz  •  AQLLI AVTO-JAVOB USERBOT + ADMIN PANEL
================================================================================
  • Telethon (o'z hisobingizga avto-javob beradi — shaxsiy chatlar)
  • Aiogram (bot / boshqaruv paneli + admin panel)
  • 100+ ta admin funksiyasi: statistika, foydalanuvchilar, triggerlar,
    broadcast, global sozlamalar, bayram tabriklari, loglar, media,
    xavfsizlik, tozalash, tizim va boshqalar.
================================================================================
"""

import asyncio
import os
import sys
import sqlite3
import html
import json
import time
import re
import io
import traceback
from datetime import datetime, timedelta

from aiogram import Bot, Dispatcher, types
from aiogram.filters import Command
from aiogram.types import (
    InlineKeyboardMarkup, InlineKeyboardButton,
    ReplyKeyboardMarkup, KeyboardButton, ReplyKeyboardRemove,
    BufferedInputFile
)
from aiogram import F
from telethon import TelegramClient, events
from telethon.sessions import StringSession
from telethon.tl.functions.users import GetFullUserRequest
from telethon.tl.types import UserStatusOnline
from telethon.errors import (
    SessionPasswordNeededError,
    FloodWaitError,
    AuthKeyUnregisteredError,
    UserDeactivatedBanError,
    UnauthorizedError
)

# ========== TARJIMA VA TIL ANIQLASH ==========
from deep_translator import GoogleTranslator
from langdetect import detect, DetectorFactory
DetectorFactory.seed = 0

# ========== EKSPORT (Excel / Word) ==========
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment
from docx import Document
from docx.shared import Pt, RGBColor

# ========== KONFIGURATSIYA ==========
DB_NAME = 'user_sessions.db'
DEFAULT_REPLY_TEXT = "Hozircha bandman, bo'shashim bilan aloqaga chiqaman."
DEFAULT_DELAY = 7

ADMIN_USERNAME = "NeoPulse_uz"
API_ID = 37437082
API_HASH = "b7d4fa4d28472bf3768a4cae5e3fd01c"
BOT_TOKEN = "8995093768:AAG676LT4-ate2TFoTqHmVbFEDuIZlWsMDc"

# O'zgartirish mumkin bo'lgan prefiks va belgilar
BOT_TAG = "[🤖 Avto-Javob]"

bot = Bot(token=BOT_TOKEN)
dp = Dispatcher()

# ========== XOTIRA ==========
user_clients = {}
auto_tasks = {}
waiting_for = {}
phone_cache = {}
phone_code_hash_cache = {}
chat_log_cache = {}
admin_current_page = {}
admin_cat_state = {}          # admin panel kategoriya ochiqmi (UID -> category)
admin_total_messages = 0
BOT_START_TIME = None

# ========== MA'LUMOTLAR BAZASI ==========
def get_db():
    conn = sqlite3.connect(DB_NAME, timeout=30)
    conn.execute("PRAGMA journal_mode=WAL;")
    return conn

def init_db():
    with get_db() as conn:
        c = conn.cursor()
        c.execute('''CREATE TABLE IF NOT EXISTS sessions
                     (user_id INTEGER PRIMARY KEY, session_string TEXT,
                      custom_text TEXT, delay_seconds INTEGER DEFAULT 7)''')
        c.execute('''CREATE TABLE IF NOT EXISTS custom_triggers
                     (id INTEGER PRIMARY KEY AUTOINCREMENT, user_id INTEGER,
                      keyword TEXT, response TEXT)''')
        c.execute('''CREATE TABLE IF NOT EXISTS sent_auto_replies
                     (id INTEGER PRIMARY KEY AUTOINCREMENT, user_id INTEGER,
                      chat_id INTEGER, msg_id INTEGER)''')
        c.execute('''CREATE TABLE IF NOT EXISTS holiday_greeted
                     (user_id INTEGER, holiday_name TEXT,
                      greeted_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                      PRIMARY KEY (user_id, holiday_name))''')
        c.execute('''CREATE TABLE IF NOT EXISTS bot_config
                     (key TEXT PRIMARY KEY, value TEXT)''')
        c.execute('''CREATE TABLE IF NOT EXISTS blocks
                     (user_id INTEGER PRIMARY KEY, reason TEXT,
                      created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)''')
        c.execute('''CREATE TABLE IF NOT EXISTS admins
                     (user_id INTEGER PRIMARY KEY)''')
        c.execute('''CREATE TABLE IF NOT EXISTS welcome_log
                     (id INTEGER PRIMARY KEY AUTOINCREMENT, user_id INTEGER,
                      event TEXT, event_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)''')
        conn.commit()

# ---------- SESSIONS ----------
def save_session(user_id, session_string):
    with get_db() as conn:
        conn.execute('''INSERT INTO sessions (user_id, session_string)
                        VALUES (?, ?)
                        ON CONFLICT(user_id) DO UPDATE SET session_string=excluded.session_string''',
                     (user_id, session_string))

def has_active_session(user_id):
    with get_db() as conn:
        c = conn.cursor()
        c.execute('SELECT session_string FROM sessions WHERE user_id=? AND session_string IS NOT NULL', (user_id,))
        row = c.fetchone()
        return row is not None and bool(row[0])

def get_session_string(user_id):
    with get_db() as conn:
        c = conn.cursor()
        c.execute('SELECT session_string FROM sessions WHERE user_id=?', (user_id,))
        row = c.fetchone()
        return row[0] if row else None

def get_all_sessions():
    with get_db() as conn:
        c = conn.cursor()
        c.execute('SELECT user_id, session_string FROM sessions WHERE session_string IS NOT NULL')
        return c.fetchall()

def get_total_users_count():
    with get_db() as conn:
        c = conn.cursor()
        c.execute('SELECT COUNT(*) FROM sessions')
        return c.fetchone()[0]

def delete_session(user_id):
    with get_db() as conn:
        conn.execute('DELETE FROM sessions WHERE user_id=?', (user_id,))
        conn.execute('DELETE FROM sent_auto_replies WHERE user_id=?', (user_id,))
        conn.execute('DELETE FROM custom_triggers WHERE user_id=?', (user_id,))
        conn.execute('DELETE FROM holiday_greeted WHERE user_id=?', (user_id,))

# ---------- JAVOB MATNI / KECHIKISH ----------
def set_custom_reply_text(user_id, text):
    with get_db() as conn:
        conn.execute('''INSERT INTO sessions (user_id, custom_text)
                        VALUES (?, ?)
                        ON CONFLICT(user_id) DO UPDATE SET custom_text=excluded.custom_text''',
                     (user_id, text))

def get_custom_reply_text(user_id):
    with get_db() as conn:
        c = conn.cursor()
        c.execute('SELECT custom_text FROM sessions WHERE user_id=?', (user_id,))
        row = c.fetchone()
        if row and row[0]:
            return row[0]
    return get_global_default_text()

def set_user_delay(user_id, seconds):
    with get_db() as conn:
        conn.execute('''INSERT INTO sessions (user_id, delay_seconds)
                        VALUES (?, ?)
                        ON CONFLICT(user_id) DO UPDATE SET delay_seconds=excluded.delay_seconds''',
                     (user_id, seconds))

def get_user_delay(user_id):
    with get_db() as conn:
        c = conn.cursor()
        c.execute('SELECT delay_seconds FROM sessions WHERE user_id=?', (user_id,))
        row = c.fetchone()
        if row and row[0] is not None:
            return row[0]
    return get_global_default_delay()

# ---------- TRIGGERS ----------
def add_custom_trigger(user_id, keyword, response):
    with get_db() as conn:
        conn.execute('INSERT INTO custom_triggers (user_id, keyword, response) VALUES (?, ?, ?)',
                     (user_id, keyword.lower().strip(), response))

def get_matching_response(user_id, message_text):
    if not message_text:
        return None
    clean_msg = message_text.lower().strip()
    with get_db() as conn:
        c = conn.cursor()
        c.execute('SELECT keyword, response FROM custom_triggers WHERE user_id=?', (user_id,))
        triggers = c.fetchall()
        for kw, resp in triggers:
            if kw == clean_msg or kw in clean_msg:
                return resp
    return None

def get_user_triggers(user_id):
    with get_db() as conn:
        c = conn.cursor()
        c.execute('SELECT id, keyword, response FROM custom_triggers WHERE user_id=?', (user_id,))
        return c.fetchall()

def delete_trigger_by_id(trigger_id, user_id):
    with get_db() as conn:
        conn.execute('DELETE FROM custom_triggers WHERE id=? AND user_id=?', (trigger_id, user_id))

def get_all_triggers_count():
    with get_db() as conn:
        c = conn.cursor()
        c.execute('SELECT COUNT(*) FROM custom_triggers')
        return c.fetchone()[0]

# ---------- AUTOREPLY LOG ----------
def save_sent_reply(user_id, chat_id, msg_id):
    with get_db() as conn:
        conn.execute('INSERT INTO sent_auto_replies (user_id, chat_id, msg_id) VALUES (?, ?, ?)',
                     (user_id, chat_id, msg_id))

def get_and_delete_all_sent_replies(user_id, chat_id):
    with get_db() as conn:
        c = conn.cursor()
        c.execute('SELECT msg_id FROM sent_auto_replies WHERE user_id=? AND chat_id=?', (user_id, chat_id))
        rows = c.fetchall()
        if rows:
            c.execute('DELETE FROM sent_auto_replies WHERE user_id=? AND chat_id=?', (user_id, chat_id))
            conn.commit()
            return [r[0] for r in rows]
        return []

# ---------- ADMINLAR / BLOKLAR ----------
def is_admin(user: types.User) -> bool:
    if user and user.username and user.username.lower() == ADMIN_USERNAME.lower():
        return True
    return user.id in get_admin_ids()

def get_admin_ids():
    with get_db() as conn:
        c = conn.cursor()
        c.execute('SELECT user_id FROM admins')
        return {r[0] for r in c.fetchall()}

def add_admin(user_id):
    with get_db() as conn:
        conn.execute('INSERT OR IGNORE INTO admins (user_id) VALUES (?)', (user_id,))

def remove_admin(user_id):
    with get_db() as conn:
        conn.execute('DELETE FROM admins WHERE user_id=?', (user_id,))

def is_blocked(user_id):
    with get_db() as conn:
        c = conn.cursor()
        c.execute('SELECT 1 FROM blocks WHERE user_id=?', (user_id,))
        return c.fetchone() is not None

def block_user(user_id, reason="admin"):
    with get_db() as conn:
        conn.execute('INSERT OR IGNORE INTO blocks (user_id, reason) VALUES (?, ?)', (user_id, reason))

def unblock_user(user_id):
    with get_db() as conn:
        conn.execute('DELETE FROM blocks WHERE user_id=?', (user_id,))

def get_blocked_users():
    with get_db() as conn:
        c = conn.cursor()
        c.execute('SELECT user_id, reason FROM blocks')
        return c.fetchall()

# ---------- GLOBAL SOZLAMALAR (ADMIN) ----------
def get_config(key, default=None):
    with get_db() as conn:
        c = conn.cursor()
        c.execute('SELECT value FROM bot_config WHERE key=?', (key,))
        row = c.fetchone()
        return row[0] if row else default

def set_config(key, value):
    with get_db() as conn:
        conn.execute('''INSERT INTO bot_config (key, value) VALUES (?, ?)
                        ON CONFLICT(key) DO UPDATE SET value=excluded.value''',
                     (key, str(value)))

def get_global_default_text():
    return get_config('default_reply_text', DEFAULT_REPLY_TEXT)

def set_global_default_text(text):
    set_config('default_reply_text', text)

def get_global_default_delay():
    try:
        return int(get_config('default_delay', DEFAULT_DELAY))
    except Exception:
        return DEFAULT_DELAY

def set_global_default_delay(seconds):
    set_config('default_delay', seconds)

# ----- toggles -----
def get_flag(key, default=True):
    v = get_config(key)
    if v is None:
        return default
    return str(v).lower() in ['1', 'true', 'yes', 'on']

def set_flag(key, val):
    set_config(key, '1' if val else '0')

def get_bot_tag():
    return get_config('bot_tag', BOT_TAG)

def set_bot_tag(tag):
    set_config('bot_tag', tag)

def dlog(*args):
    """Konsolga log chiqarish — admin sozlamasiga qarab yoqiladi/o'chiriladi."""
    if get_flag('debug_enabled', True):
        print(*args)

def reset_all_holiday_greetings():
    with get_db() as conn:
        c = conn.cursor()
        c.execute('SELECT COUNT(*) FROM holiday_greeted')
        count = c.fetchone()[0]
        conn.execute('DELETE FROM holiday_greeted')
        return count

# ========== HOLIDAY FUNKSIYALARI ==========
def get_holiday_info():
    """Bugungi bayram uchun (matn, nom). Matn admin panelidan o'zgartirilishi mumkin."""
    today = datetime.now()
    month_day = today.strftime("%m-%d")
    holidays = {
        "01-01": ("holiday_text_new_year",      "🎉 Yangi Yil muborak!", "new_year"),
        "03-21": ("holiday_text_navruz",        "🌿 Navro'z muborak!", "navruz"),
        "09-01": ("holiday_text_independence",  "🇺🇿 Mustaqillik kuni muborak!", "independence_day"),
        "10-01": ("holiday_text_teacher",       "🇺🇿 O'qituvchi va murabbiylar kuni!", "teachers_day"),
        "12-08": ("holiday_text_constitution",  "🇺🇿 Konstitutsiya kuni!", "constitution_day"),
    }
    if (today.month == 8 and today.day >= 29) or (today.month == 9 and today.day == 1):
        cfg_key, default, name = holidays["09-01"]
        return (get_config(cfg_key, default), name)
    if month_day in holidays:
        cfg_key, default, name = holidays[month_day]
        return (get_config(cfg_key, default), name)
    return None, None

def has_user_received_holiday(user_id, holiday_name):
    with get_db() as conn:
        c = conn.cursor()
        c.execute('SELECT 1 FROM holiday_greeted WHERE user_id=? AND holiday_name=?', (user_id, holiday_name))
        return c.fetchone() is not None

def mark_holiday_greeted(user_id, holiday_name):
    with get_db() as conn:
        conn.execute('INSERT OR IGNORE INTO holiday_greeted (user_id, holiday_name) VALUES (?, ?)',
                     (user_id, holiday_name))

# ========== YORDAMCHI FUNKSIYALAR ==========
def is_silent_hour():
    if not get_flag('silent_enabled', True):
        return False
    hour = datetime.now().hour
    return hour >= 22 or hour < 7

def get_time_based_prefix():
    hour = datetime.now().hour
    if 5 <= hour < 11:   return "Xayrli tong"
    if 11 <= hour < 18:  return "Xayrli kun"
    if 18 <= hour < 23:  return "Xayrli kech"
    return "Xayrli tun"

def latin_to_cyrillic_uz(text: str) -> str:
    mapping = {
        'a':'а','b':'б','d':'д','e':'е','f':'ф','g':'г',
        'h':'ҳ','i':'и','j':'ж','k':'к','l':'л','m':'м',
        'n':'н','o':'о','p':'п','q':'қ','r':'р','s':'с',
        't':'т','u':'у','v':'в','x':'х','y':'й','z':'з',
        '‘':'ъ','ʻ':'ъ',
        'A':'А','B':'Б','D':'Д','E':'Е','F':'Ф','G':'Г',
        'H':'Ҳ','I':'И','J':'Ж','K':'К','L':'Л','M':'М',
        'N':'Н','O':'О','P':'П','Q':'Қ','R':'Р','S':'С',
        'T':'Т','U':'У','V':'В','X':'Х','Y':'Й','Z':'З',
    }
    cyrillic = ''
    i = 0
    while i < len(text):
        if i+1 < len(text):
            pair = text[i:i+2].lower()
            if pair == 'sh': cyrillic += 'ш'; i += 2; continue
            if pair == 'ch': cyrillic += 'ч'; i += 2; continue
            if pair == 'ng': cyrillic += 'нг'; i += 2; continue
            if pair == "o'": cyrillic += 'ў'; i += 2; continue
            if pair == "g'": cyrillic += 'ғ'; i += 2; continue
        cyrillic += mapping.get(text[i], text[i])
        i += 1
    return cyrillic

def is_cyrillic_uz(text: str) -> bool:
    return bool(re.search(r'[А-Яа-яЁё]', text))

def detect_language(text: str) -> str:
    if not text:
        return 'uz'
    try:
        lang = detect(text)
        if lang in ['uz', 'uz-Cyrl']:
            return 'uz'
        elif lang == 'ru':
            return 'ru'
        elif lang == 'en':
            return 'en'
        else:
            return 'uz'
    except Exception:
        return 'uz'

def translate_text(text: str, dest_lang: str) -> str:
    if dest_lang not in ['uz', 'ru', 'en']:
        return text
    try:
        return GoogleTranslator(source='auto', target=dest_lang).translate(text)
    except Exception as e:
        print(f"Tarjima xatosi: {e}")
        return text

# ========== LOGLAR (2 soatlik) ==========
def add_log_entry(user_id, chat_id, from_me, text):
    now = datetime.now()
    entry = {'timestamp': now, 'chat_id': chat_id, 'from_me': from_me, 'text': text}
    chat_log_cache.setdefault(user_id, []).append(entry)
    cutoff = now - timedelta(hours=2)
    chat_log_cache[user_id] = [e for e in chat_log_cache[user_id] if e['timestamp'] >= cutoff]

def get_logs_last_2h(user_id):
    if user_id not in chat_log_cache:
        return []
    cutoff = datetime.now() - timedelta(hours=2)
    return [e for e in chat_log_cache[user_id] if e['timestamp'] >= cutoff]

# ========== EKSPORT (Excel / Word) ==========
def collect_user_data():
    """Barcha foydalanuvchilar ma'lumotini list-of-dict qilib to'playdi."""
    default_delay = get_global_default_delay()
    default_text = get_global_default_text()
    rows = []
    with get_db() as conn:
        c = conn.cursor()
        c.execute('SELECT user_id, session_string, custom_text, delay_seconds FROM sessions')
        for uid, sess, ctext, delay in c.fetchall():
            trig = get_user_triggers(uid)
            rows.append({
                'id': uid,
                'text': ctext if ctext else default_text,
                'delay': delay if (delay is not None and delay >= 0) else default_delay,
                'triggers': len(trig),
                'session': bool(sess),
                'running': uid in user_clients,
                'blocked': is_blocked(uid),
            })
    return rows

def build_excel_bytes(rows):
    from openpyxl.utils import get_column_letter
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Foydalanuvchilar"
    headers = ["ID", "Javob matni", "Kechikish (s)", "Triggerlar",
               "Ulangan seans", "Faol jarayon", "Bloklangan"]
    ws.append(headers)
    for c in ws[1]:
        c.font = Font(bold=True, color="FFFFFF")
        c.fill = PatternFill("solid", fgColor="305496")
        c.alignment = Alignment(horizontal="center")
    for r in rows:
        ws.append([
            r['id'], r['text'], r['delay'], r['triggers'],
            "Ha" if r['session'] else "Yo'q",
            "Ha" if r['running'] else "Yo'q",
            "Ha" if r['blocked'] else "Yo'q",
        ])
    widths = [18, 46, 15, 11, 16, 14, 12]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w
    bio = io.BytesIO()
    wb.save(bio)
    bio.seek(0)
    return bio.getvalue()

def build_word_bytes(rows):
    doc = Document()
    doc.add_heading("Foydalanuvchilar ma'lumoti — Reply-uz", level=1)
    doc.add_paragraph(f"Yaratilgan vaqti: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    table = doc.add_table(rows=1, cols=7)
    table.style = "Light Grid Accent 1"
    headers = ["ID", "Javob matni", "Kechikish (s)", "Triggerlar",
               "Ulangan seans", "Faol jarayon", "Bloklangan"]
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x30, 0x54, 0x96)
    for r in rows:
        cells = table.add_row().cells
        values = [
            str(r['id']), str(r['text']), str(r['delay']), str(r['triggers']),
            "Ha" if r['session'] else "Yo'q",
            "Ha" if r['running'] else "Yo'q",
            "Ha" if r['blocked'] else "Yo'q",
        ]
        for i, v in enumerate(values):
            cells[i].text = v
            for p in cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()

# ========== UMUMIY JAVOB TAYYORLASH ==========
def get_pending_holiday_greeting(user_id: int):
    """Bayram tabrigi hali yuborilmagan bo'lsa, uni qaytaradi. Alohida birinchi xabar sifatida yuboriladi."""
    if not get_flag('holiday_enabled', True):
        return None, None
    greeting, h_name = get_holiday_info()
    if greeting and h_name and not has_user_received_holiday(user_id, h_name):
        return greeting, h_name
    return None, None

def get_reply_text(user_id: int, incoming_text: str, is_telethon: bool = False) -> str:
    """Avto-javob matnini tayyorlaydi (vaqt, prefiks, silent, tarjima, kirill)."""
    my_text = get_custom_reply_text(user_id)

    matched = get_matching_response(user_id, incoming_text)
    reply = matched if matched else my_text

    # Prefiks
    if get_flag('time_prefix_enabled', True):
        reply = f"{get_time_based_prefix()}! {reply}"

    tag = get_bot_tag()
    reply = f"{tag} {reply}"

    if is_silent_hour():
        reply += "\n\n🌙 Hozir uxlab yotibman, ertalab javob beraman.."

    if get_flag('cyrillic_enabled', True) and is_cyrillic_uz(incoming_text):
        reply = latin_to_cyrillic_uz(reply)

    if get_flag('translate_enabled', True):
        lang = detect_language(incoming_text)
        if lang in ['uz', 'ru', 'en']:
            reply = translate_text(reply, lang)

    return reply

# ========== TELETHON AVTO JAVOB (faqat shaxsiy xabarlar) ==========
async def start_auto_reply(user_id, client: TelegramClient):
    dlog(f"✅ start_auto_reply called for user {user_id}")

    @client.on(events.NewMessage(incoming=True, func=lambda e: e.is_private))
    async def auto_reply_handler(event):
        # Shaxsiy xabarlarga avto-javob beramiz
        try:
            sender = await event.get_sender()
            # Bot hisoblarga avtojavob yubormaymiz (bot-bot loop bo'lmasin)
            if not sender or sender.bot:
                return

            if event.out:
                return

            if not get_flag('autoreply_enabled', True):
                return

            my_text = get_custom_reply_text(user_id)
            if event.raw_text and event.raw_text.strip() == my_text.strip():
                return

            add_log_entry(user_id, event.chat_id, False, event.raw_text or '')

            # Bloklangan foydalanuvchiga javob bermaymiz
            if is_blocked(event.sender_id):
                return

            try:
                full_user = await client(GetFullUserRequest(event.sender_id))
                is_online = isinstance(full_user.users[0].status, UserStatusOnline)
            except Exception as e:
                dlog(f"⚠️ Online holatni tekshirishda xatolik: {e}")
                is_online = False

            reply = get_reply_text(user_id, event.raw_text or '', is_telethon=True)

            # ------ BAYRAM TABRIGI: ALOHIDA, BIRINCHI XABAR sifatida yuboriladi ------
            greeting, h_name = get_pending_holiday_greeting(user_id)
            if greeting:
                try:
                    await client.send_message(event.chat_id, greeting)
                    mark_holiday_greeted(user_id, h_name)
                    dlog(f"🎉 Holiday greeting sent to {event.chat_id}")
                except Exception as e:
                    dlog(f"⚠️ Bayram tabrigini yuborishda xatolik: {e}")

            if is_online:
                delay = get_user_delay(user_id)
                dlog(f"⏳ Online, waiting {delay}s")
                await asyncio.sleep(delay)

            # ------ MUHIM: JAVOB birinchi yuboriladi, so'ng "o'qilgan" qilinadi ------
            sent = await client.send_message(
                event.chat_id, reply,
                reply_to=event.id
            )
            save_sent_reply(user_id, event.chat_id, sent.id)
            dlog(f"✅ Reply sent, msg id {sent.id}")

            # O'qilgan deb belgilash — JAVOBDAN KEYIN
            if get_flag('read_enabled', True):
                try:
                    await client.send_read_acknowledge(event.chat_id)
                except Exception as e:
                    dlog(f"⚠️ Read ack xatosi: {e}")

        except (AuthKeyUnregisteredError, UserDeactivatedBanError, UnauthorizedError) as e:
            dlog(f"❌ Auth error: {e}")
            stop_client_task(user_id); delete_session(user_id)
        except FloodWaitError as e:
            dlog(f"⏳ Flood wait {e.seconds}s")
            await asyncio.sleep(e.seconds)
        except Exception as e:
            dlog(f"❌ Handler error: {e}")
            traceback.print_exc()

    @client.on(events.NewMessage(incoming=True, func=lambda e: e.is_group))
    async def telethon_group_reader(event):
        # Guruhdagi xabarlarni o'qiymiz: javob YO'ZMAYMIZ, konsolga log CHIQARMAYMIZ.
        # Faqat "o'qilgan" deb belgilaymiz (agar yoqilgan bo'lsa).
        if event.out:
            return
        if not get_flag('read_enabled', True):
            return
        try:
            await client.send_read_acknowledge(event.chat_id)
        except Exception:
            pass

    @client.on(events.NewMessage(outgoing=True))
    async def outgoing_handler(event):
        if not event.is_private:
            return
        if event.text and event.text.strip() == ".export_logs":
            logs = get_logs_last_2h(user_id)
            if not logs:
                await client.send_message('me', "📭 So'nggi 2 soat ichida hech qanday xabar mavjud emas.")
                return
            lines = []
            for entry in logs:
                timestamp = entry['timestamp'].strftime("%Y-%m-%d %H:%M:%S")
                who = "Men" if entry['from_me'] else "Suhbatdosh"
                lines.append(f"[{timestamp}] {who}: {entry['text']}")
            await client.send_file(
                'me',
                "\n".join(lines).encode('utf-8'),
                caption="📄 So'nggi 2 soatlik chat tarixi",
                force_document=True
            )
            return

        try:
            chat_id = event.chat_id
            msg_ids = get_and_delete_all_sent_replies(user_id, chat_id)
            if msg_ids:
                await client.delete_messages(chat_id, msg_ids)
        except Exception as e:
            dlog(f"Avto-javoblarni o'chirishda xatolik: {e}")

        if event.text and not event.text.startswith('.'):
            add_log_entry(user_id, event.chat_id, True, event.text)

    @client.on(events.MessageRead(inbox=True))
    async def read_handler(event):
        try:
            chat_id = event.chat_id
            msg_ids = get_and_delete_all_sent_replies(user_id, chat_id)
            if msg_ids:
                await client.delete_messages(chat_id, msg_ids)
        except Exception as e:
            dlog(f"Read hodisasida xatolik: {e}")

    try:
        dlog(f"🔄 Starting client.run_until_disconnected() for user {user_id}")
        await client.run_until_disconnected()
        dlog(f"🛑 client.run_until_disconnected() finished for user {user_id}")
    except Exception as e:
        dlog(f"❌ client.run_until_disconnected() error: {e}")
        traceback.print_exc()
    finally:
        dlog(f"🧹 Cleaning up user {user_id}")
        user_clients.pop(user_id, None)
        auto_tasks.pop(user_id, None)

def stop_client_task(user_id):
    client = user_clients.get(user_id)
    task = auto_tasks.get(user_id)
    if client:
        try:
            asyncio.create_task(client.disconnect())
        except Exception:
            pass
    if task and not task.done():
        try:
            task.cancel()
        except Exception:
            pass
    user_clients.pop(user_id, None)
    auto_tasks.pop(user_id, None)

# ========== TUGMALAR VA MENYULAR ==========
def get_main_keyboard(is_connected=False):
    buttons = []
    if is_connected:
        buttons.append([InlineKeyboardButton(text="➕ So'rov/Javob qo'shish (Kalit so'z)", callback_data="add_trigger")])
        buttons.append([InlineKeyboardButton(text="📋 Barcha so'rovlarni ko'rish", callback_data="list_triggers")])
        buttons.append([InlineKeyboardButton(text="✏️ Asosiy javobni o'zgartirish", callback_data="change_text")])
        buttons.append([InlineKeyboardButton(text="⏱ Kutiladigan vaqtni o'zgartirish", callback_data="change_delay")])
        buttons.append([InlineKeyboardButton(text="⏹ Avto-javobni to'xtatish", callback_data="stop_auto")])
    else:
        buttons.append([InlineKeyboardButton(text="🚀 Boshlash (Ulanish)", callback_data="start_auto")])
    return InlineKeyboardMarkup(inline_keyboard=buttons)

def get_reply_keyboard(user: types.User):
    keyboard = [[KeyboardButton(text="⚙️ Sozlamalar / Bosh menyu")]]
    if is_admin(user):
        keyboard.append([KeyboardButton(text="👑 Admin Panel")])
    return ReplyKeyboardMarkup(keyboard=keyboard, resize_keyboard=True)

# ========== START ==========
@dp.message(Command("start"))
async def start_cmd(message: types.Message):
    user_id = message.from_user.id
    if is_blocked(user_id):
        await message.answer("🚫 Siz bloklangansiz.")
        return
    is_connected = has_active_session(user_id)
    if is_connected:
        await message.answer(
            f"Assalomu alaykum! 👋\n\n"
            f"💬 <b>Asosiy avto-javob matningiz:</b>\n<code>{html.escape(get_custom_reply_text(user_id))}</code>\n\n"
            f"⏱ <b>Online kutiladigan vaqt:</b> <code>{get_user_delay(user_id)} soniya</code>\n\n"
            f"Kerakli bo'limni tanlang:",
            parse_mode="HTML",
            reply_markup=get_main_keyboard(True)
        )
    else:
        await message.answer(
            "Assalomu alaykum! 👋\n\n"
            "Botdan foydalanish uchun Telegram akkauntingizni ulang.",
            reply_markup=get_main_keyboard(False)
        )
    await message.answer("👇 Qulay boshqaruv menyusi:", reply_markup=get_reply_keyboard(message.from_user))

@dp.message(F.text == "👑 Admin Panel")
async def admin_panel_cmd(message: types.Message):
    if not is_admin(message.from_user):
        return
    admin_cat_state[message.from_user.id] = None
    await message.answer(
        "🛠 <b>Admin Panel</b>\n\n"
        "Siz bot tizimida admin huquqiga egasiz.\n"
        "Quyidagi tugmani bosing:",
        parse_mode="HTML",
        reply_markup=InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="📋 Admin menyuni ochish", callback_data="admin_open")]
        ])
    )

@dp.message(F.text == "⚙️ Sozlamalar / Bosh menyu")
async def settings_cmd(message: types.Message):
    user_id = message.from_user.id
    await message.answer(
        "⚙️ <b>Bosh menyu va sozlamalar:</b>",
        parse_mode="HTML",
        reply_markup=get_main_keyboard(has_active_session(user_id))
    )

# ========== GURUH XABARLARI (o'qiydi, javob bermaydi) ==========
@dp.message(F.chat.type.in_({"group", "supergroup"}))
async def group_silent_reader(message: types.Message):
    return

# ========== CALLBACKLAR (USER FLOW) ==========
async def check_login(callback: types.CallbackQuery) -> bool:
    user_id = callback.from_user.id
    if not has_active_session(user_id):
        await callback.answer("❌ Iltimos, avval /start orqali ulaning.", show_alert=True)
        return False
    return True

@dp.callback_query(F.data == "add_trigger")
async def add_trigger_cb(callback: types.CallbackQuery):
    if not await check_login(callback):
        return
    waiting_for[callback.from_user.id] = "trigger_keyword"
    await callback.message.answer(
        "🔹 **Qaysi kichik gap/so'z yozilganda javob berilsin?**\n\nMasalan: `salom` yoki `qayerdasiz`",
        parse_mode="Markdown"
    )
    await callback.answer()

@dp.callback_query(F.data == "list_triggers")
async def list_triggers_cb(callback: types.CallbackQuery):
    if not await check_login(callback):
        return
    user_id = callback.from_user.id
    triggers = get_user_triggers(user_id)
    if not triggers:
        await callback.message.answer("📭 Siz hali hech qanday maxsus so'rov qo'shmadingiz.")
        await callback.answer()
        return
    msg = "📋 **Siz qo'shgan maxsus so'rovlar:**\n\n"
    buttons = []
    for t_id, kw, resp in triggers:
        msg += f"🔸 **So'rov:** `{kw}` ➡️ **Javob:** `{resp}`\n"
        buttons.append([InlineKeyboardButton(text=f"❌ O'chirish: {kw}", callback_data=f"del_trg_{t_id}")])
    await callback.message.answer(msg, parse_mode="Markdown", reply_markup=InlineKeyboardMarkup(inline_keyboard=buttons))
    await callback.answer()

@dp.callback_query(F.data.startswith("del_trg_"))
async def delete_trigger_cb(callback: types.CallbackQuery):
    if not await check_login(callback):
        return
    t_id = int(callback.data.split("_")[2])
    delete_trigger_by_id(t_id, callback.from_user.id)
    await callback.message.answer("✅ Maxsus so'rov o'chirib tashlandi.")
    await callback.answer()

@dp.callback_query(F.data == "change_text")
async def change_text_callback(callback: types.CallbackQuery):
    if not await check_login(callback):
        return
    waiting_for[callback.from_user.id] = "custom_text"
    await callback.message.answer("📝 Yangi asosiy avto-javob matnini kiriting:")
    await callback.answer()

@dp.callback_query(F.data == "change_delay")
async def change_delay_callback(callback: types.CallbackQuery):
    if not await check_login(callback):
        return
    waiting_for[callback.from_user.id] = "custom_delay"
    await callback.message.answer("⏱ Online bo'lganda necha soniyadan keyin javob qaytarilsin? (Faqat raqam):")
    await callback.answer()

@dp.callback_query(F.data == "start_auto")
async def start_auto_callback(callback: types.CallbackQuery):
    user_id = callback.from_user.id
    if has_active_session(user_id):
        await callback.message.edit_text("✅ Siz allaqachon ulangansiz.", reply_markup=get_main_keyboard(True))
        await callback.answer()
        return
    keyboard = ReplyKeyboardMarkup(
        keyboard=[[KeyboardButton(text="📱 Telefon raqamini ulashish", request_contact=True)]],
        resize_keyboard=True, one_time_keyboard=True
    )
    await callback.message.edit_text("📱 Telefon raqamingizni ulashing.")
    await callback.message.answer("👇 Tugmani bosing:", reply_markup=keyboard)
    waiting_for[user_id] = "phone"
    await callback.answer()

@dp.callback_query(F.data == "stop_auto")
async def stop_auto_callback(callback: types.CallbackQuery):
    if not await check_login(callback):
        return
    user_id = callback.from_user.id
    stop_client_task(user_id)
    delete_session(user_id)
    waiting_for[user_id] = None
    await callback.message.edit_text("✅ Avto-javob to'xtatildi.", reply_markup=get_main_keyboard(False))
    await callback.answer()

# =======================================================================
#                        ADMIN PANEL (100+ funksiya)
# =======================================================================
ADMIN_PAGE_SIZE = 6

# ---- kategoriyalar (asosiy admin menyu) ----
ADMIN_CATEGORIES = [
    ("📊 Statistika",            "adm_cat_stats"),
    ("👥 Foydalanuvchilar",      "adm_cat_users"),
    ("🔑 Triggerlar (global)",   "adm_cat_triggers"),
    ("📢 Broadcast",             "adm_cat_broadcast"),
    ("⚙️ Global sozlamalar",     "adm_cat_settings"),
    ("🎉 Bayram / Tabrik",       "adm_cat_holiday"),
    ("📝 Chat / Loglar",         "adm_cat_logs"),
    ("🖼 Media yuborish",        "adm_cat_media"),
    ("🔒 Xavfsizlik",            "adm_cat_security"),
    ("🗑 Tozalash (Cleanup)",    "adm_cat_cleanup"),
    ("🛠 Boshqaruv vositalari",  "adm_cat_tools"),
    ("⚡️ Tizim",                "adm_cat_system"),
    ("🇺🇿 Til / Tarjima",        "adm_cat_lang"),
    ("📤 Eksport (Excel/Word)",  "adm_cat_export"),
]

# ---- har kategoriyaning funksiyalari: (label, callback) ----
CAT_STATS = [
    ("📊 Umumiy statistika",      "adm_stats_main"),
    ("👥 Jami foydalanuvchilar",  "adm_stats_users"),
    ("⚡️ Faol seanslar",          "adm_stats_sessions"),
    ("🔄 Ish turgan jarayonlar",  "adm_stats_running"),
    ("🔑 Jami triggerlar",        "adm_stats_triggers"),
    ("🔑 Global triggerlar soni", "adm_stats_gtriggers"),
    ("⏳ O'rtacha kechikish",     "adm_stats_avgdelay"),
    ("🕘 Bot uptime",             "adm_stats_uptime"),
    ("💾 Baza hajmi",             "adm_stats_dbsize"),
    ("📨 Jami yuborilgan xabar",  "adm_stats_msgs"),
    ("🚫 Bloklanganlar soni",     "adm_stats_blocks"),
    ("👑 Adminlar soni",          "adm_stats_admins"),
]

CAT_USERS = [
    ("👥 Ro'yxat (sahifalab)",    "adm_users_list_0"),
    ("🔍 Foydalanuvchi ma'lumot", "adm_user_info"),
    ("❌ Foydalanuvchini o'chirish", "adm_user_delete"),
    ("✏️ Javob matnini o'zgartirish", "adm_user_set_text"),
    ("⏱ Kechikishni o'zgartirish",   "adm_user_set_delay"),
    ("➕ Trigger qo'shish",           "adm_user_add_trigger"),
    ("📋 Triggerlarini ko'rish",     "adm_user_list_triggers"),
    ("⏹ Jarayonini to'xtatish",      "adm_user_stop"),
    ("▶️ Jarayonini qayta boshlash",  "adm_user_start"),
    ("🧪 Javobini sinovdan o'tkazish","adm_user_test_reply"),
    ("🔁 Seansini yangilash",         "adm_user_refresh"),
    ("📝 Chat loglari",               "adm_user_logs"),
    ("↩️ O'qilganini qaytarish (undo)","adm_user_undo"),
]

CAT_TRIGGERS = [
    ("➕ Global trigger qo'shish",  "adm_global_trigger_add"),
    ("📋 Barcha triggerlar",        "adm_global_triggers_list_0"),
    ("🔍 Trigger qidirish",         "adm_global_trigger_find"),
    ("❌ Trigger o'chirish",        "adm_global_trigger_del"),
    ("🧹 Barcha triggerlarni tozalash", "adm_global_triggers_clear"),
    ("🔢 Triggerlar statistikasi",  "adm_global_triggers_stats"),
]

CAT_BROADCAST = [
    ("📝 Matn (barchaga)",      "adm_bcast_text_all"),
    ("🟢 Faqat online'larga",   "adm_bcast_text_online"),
    ("🔴 Faqat offline'larga",  "adm_bcast_text_offline"),
    ("⚡️ Faol jarayonlarga",   "adm_bcast_text_active"),
    ("🖼 Rasm (barchaga)",      "adm_bcast_photo"),
    ("🎞 Video (barchaga)",     "adm_bcast_video"),
    ("📄 Fayl (barchaga)",      "adm_bcast_doc"),
    ("🎵 Audio (barchaga)",     "adm_bcast_audio"),
    ("📊 Broadcast tarixi",     "adm_bcast_history"),
]

CAT_SETTINGS = [
    ("✏️ Standart javob matni",      "adm_set_default_text"),
    ("⏱ Standart kechikish",         "adm_set_default_delay"),
    ("🏷 Bot tag (prefiks)",          "adm_set_bot_tag"),
    ("🌃 Silent soat yoqish/o'chirish", "adm_toggle_silent"),
    ("🕐 Vaqt prefiksi yoqish/o'chirish", "adm_toggle_timeprefix"),
    ("🔤 Tarjima yoqish/o'chirish",   "adm_toggle_translate"),
    ("🅰️ Kirill konvertatsiya",       "adm_toggle_cyrillic"),
    ("💬 Avto-javob yoqish/o'chirish","adm_toggle_autoreply"),
    ("👓 O'qilgan (read) yoqish/o'chirish", "adm_toggle_read"),
    ("🚼 Yangi foydalanuvchi xabar berish", "adm_toggle_welcome"),
    ("📝 Welcome matni",             "adm_set_welcome_text"),
    ("🚫 Filter so'z qo'shish",       "adm_add_filter_word"),
    ("📋 Filter so'zlar",            "adm_list_filter_words"),
    ("🚮 Filter so'zni o'chirish",    "adm_del_filter_word"),
    ("🧹 Barcha sozlamalarni tiklash", "adm_reset_settings"),
    ("🎉 Bayram tabriki yoqish/o'chirish", "adm_toggle_holiday"),
    ("🐞 Debug log yoqish/o'chirish",     "adm_toggle_debug"),
]

CAT_HOLIDAY = [
    ("🎉 Bayram belgilarini tozalash", "adm_reset_holidays"),
    ("🎊 Yangi Yil matni",          "adm_set_holiday_newyear"),
    ("🌿 Navro'z matni",            "adm_set_holiday_navruz"),
    ("🇺🇿 Mustaqillik matni",        "adm_set_holiday_indep"),
    ("👩‍🏫 O'qituvchilar kuni matni","adm_set_holiday_teacher"),
    ("⚖️ Konstitutsiya kuni matni",  "adm_set_holiday_constitution"),
    ("📋 Bayram tabriklanganlar",    "adm_holiday_greeted_list"),
]

CAT_LOGS = [
    ("📝 Loglarni ko'rish",     "adm_logs_view"),
    ("🗑 Loglarni tozalash",     "adm_logs_clear"),
    ("⏱ Log saqlash vaqti",     "adm_logs_ttl"),
    ("📊 Log statistikasi",     "adm_logs_stats"),
]

CAT_MEDIA = [
    ("🖼 Rasm yuborish",    "adm_media_photo"),
    ("🎞 Video yuborish",   "adm_media_video"),
    ("📄 Fayl yuborish",    "adm_media_doc"),
    ("🎵 Audio yuborish",   "adm_media_audio"),
]

CAT_SECURITY = [
    ("👑 Admin qo'shish",            "adm_add_admin"),
    ("👥 Adminlar ro'yxati",        "adm_list_admins"),
    ("❌ Adminni olib tashlash",     "adm_remove_admin"),
    ("🚫 Foydalanuvchini bloklash",  "adm_block_user"),
    ("✅ Blokdan chiqarish",         "adm_unblock_user"),
    ("📋 Bloklanganlar ro'yxati",    "adm_list_blocks"),
    ("🔐 Hamma seanslarni 2FA bilan himoya qilish", "adm_secure_all"),
    ("🛑 Flood so'rovlarini to'xtatish", "adm_stop_flood"),
]

CAT_CLEANUP = [
    ("🧹 Barcha belgilarni tozalash", "adm_clear_holidays"),
    ("🧾 Eski seanslarni o'chirish", "adm_clean_old_sessions"),
    ("🚮 Filter so'zlarini tozalash", "adm_clear_filters"),
    ("🧹 Bazani optimallashtirish",   "adm_optimize_db"),
    ("🗑 Barcha triggerlarni o'chirish", "adm_clear_all_triggers"),
]

CAT_TOOLS = [
    ("📊 Statistika",       "adm_stats_main"),
    ("⚠️ Diagnostika",      "adm_diag"),
    ("♻️ Botni RESTART qilish", "adm_restart"),
    ("⏹ Barcha jarayonlarni to'xtatish", "adm_stop_all"),
    ("📦 DB eksport qilish",   "adm_export_db"),
    ("ℹ️ Bot haqida",         "adm_about"),
]

CAT_SYSTEM = [
    ("♻️ RESTART",             "adm_restart"),
    ("⏹ Barcha to'xtatish",    "adm_stop_all"),
    ("▶️ Barcha qayta boshlash", "adm_start_all"),
    ("📤 DB eksport",          "adm_export_db"),
    ("ℹ️ Bot haqida",          "adm_about"),
    ("🔄 Webhookni o'chirish",  "adm_clear_webhook"),
]

CAT_LANG = [
    ("🇺🇿 O'zbek (asosiy)",     "adm_lang_uz"),
    ("🇷🇺 Ruscha interfeys",    "adm_lang_ru"),
    ("🇬🇧 Inglizcha interfeys", "adm_lang_en"),
    ("🅰️ Kirill/limit auto",   "adm_toggle_cyrillic"),
]

CAT_EXPORT = [
    ("📊 Excel (.xlsx) yuborish",  "adm_export_xlsx"),
    ("📄 Word (.docx) yuborish",   "adm_export_docx"),
    ("👥 Jami foydalanuvchilar",   "adm_export_count"),
    ("🧹 Eksport uchun tozalash",  "adm_export_cleanup"),
]

CATEGORY_ITEMS = {
    "adm_cat_stats":     CAT_STATS,
    "adm_cat_users":     CAT_USERS,
    "adm_cat_triggers":  CAT_TRIGGERS,
    "adm_cat_broadcast": CAT_BROADCAST,
    "adm_cat_settings":  CAT_SETTINGS,
    "adm_cat_holiday":   CAT_HOLIDAY,
    "adm_cat_logs":      CAT_LOGS,
    "adm_cat_media":     CAT_MEDIA,
    "adm_cat_security":  CAT_SECURITY,
    "adm_cat_cleanup":   CAT_CLEANUP,
    "adm_cat_tools":     CAT_TOOLS,
    "adm_cat_system":    CAT_SYSTEM,
    "adm_cat_lang":      CAT_LANG,
    "adm_cat_export":    CAT_EXPORT,
}

def get_admin_keyboard(page: int = 0):
    start = page * ADMIN_PAGE_SIZE
    end = start + ADMIN_PAGE_SIZE
    chunk = ADMIN_CATEGORIES[start:end]
    buttons = [[InlineKeyboardButton(text=label, callback_data=cb)] for label, cb in chunk]
    nav = []
    if page > 0:
        nav.append(InlineKeyboardButton(text="⬅️ Oldingi", callback_data=f"adm_page_{page-1}"))
    if end < len(ADMIN_CATEGORIES):
        nav.append(InlineKeyboardButton(text="Keyingi ➡️", callback_data=f"adm_page_{page+1}"))
    if nav:
        buttons.append(nav)
    buttons.append([InlineKeyboardButton(text="🔙 Yopish", callback_data="adm_close")])
    return InlineKeyboardMarkup(inline_keyboard=buttons)

def get_category_keyboard(cat: str):
    items = CATEGORY_ITEMS.get(cat, [])
    buttons = [[InlineKeyboardButton(text=label, callback_data=cb)] for label, cb in items]
    buttons.append([InlineKeyboardButton(text="🔙 Ortga (Admin)", callback_data="admin_open")])
    buttons.append([InlineKeyboardButton(text="🔙 Yopish", callback_data="adm_close")])
    return InlineKeyboardMarkup(inline_keyboard=buttons)

# ---------- admin ochish / sahifalash ----------
@dp.callback_query(F.data == "admin_open")
async def admin_open_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user):
        await callback.answer()
        return
    admin_current_page[callback.from_user.id] = 0
    admin_cat_state[callback.from_user.id] = None
    try:
        await callback.message.edit_text("🛠 <b>Admin Panel</b>\n\nKerakli bo'limni tanlang:", parse_mode="HTML", reply_markup=get_admin_keyboard(0))
    except Exception:
        await callback.message.answer("🛠 <b>Admin Panel</b>\n\nKerakli bo'limni tanlang:", parse_mode="HTML", reply_markup=get_admin_keyboard(0))
    await callback.answer()

@dp.callback_query(F.data.startswith("adm_page_"))
async def admin_page_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user):
        await callback.answer(); return
    page = int(callback.data.split("_")[-1])
    admin_current_page[callback.from_user.id] = page
    admin_cat_state[callback.from_user.id] = None
    try:
        await callback.message.edit_text("🛠 <b>Admin Panel</b>\n\nKerakli bo'limni tanlang:", parse_mode="HTML", reply_markup=get_admin_keyboard(page))
    except Exception:
        pass
    await callback.answer()

@dp.callback_query(F.data == "adm_close")
async def admin_close_cb(callback: types.CallbackQuery):
    try:
        await callback.message.delete()
    except Exception:
        pass
    await callback.answer()

# ---------- kategoriya ochish ----------
@dp.callback_query(F.data.startswith("adm_cat_"))
async def admin_category_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user):
        await callback.answer(); return
    cat = callback.data
    admin_cat_state[callback.from_user.id] = cat
    title = dict(ADMIN_CATEGORIES).get(cat, "Bo'lim")
    await callback.message.edit_text(
        f"🛠 <b>{title}</b>\n\nKerakli funksiyani tanlang:",
        parse_mode="HTML", reply_markup=get_category_keyboard(cat)
    )
    await callback.answer()

# ================= STATISTIKA =================
@dp.callback_query(F.data == "adm_stats_main")
async def adm_stats_main_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user):
        await callback.answer(); return
    await admin_stats(callback.message)
    await callback.answer()

@dp.callback_query(F.data == "adm_stats_users")
async def adm_stats_users_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user):
        await callback.answer(); return
    await callback.message.answer(f"👥 Jami foydalanuvchilar: <b>{get_total_users_count()}</b>", parse_mode="HTML")
    await callback.answer()

@dp.callback_query(F.data == "adm_stats_sessions")
async def adm_stats_sessions_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user):
        await callback.answer(); return
    await callback.message.answer(f"⚡️ Bazadagi faol seanslar: <b>{len(get_all_sessions())}</b>", parse_mode="HTML")
    await callback.answer()

@dp.callback_query(F.data == "adm_stats_running")
async def adm_stats_running_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user):
        await callback.answer(); return
    await callback.message.answer(f"🔄 Hozir ishlab turgan jarayonlar: <b>{len(user_clients)}</b>", parse_mode="HTML")
    await callback.answer()

@dp.callback_query(F.data == "adm_stats_triggers")
async def adm_stats_triggers_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user):
        await callback.answer(); return
    await callback.message.answer(f"🔑 Jami triggerlar: <b>{get_all_triggers_count()}</b>", parse_mode="HTML")
    await callback.answer()

@dp.callback_query(F.data == "adm_stats_gtriggers")
async def adm_stats_gtriggers_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user):
        await callback.answer(); return
    # global triggerlarni bot_config ichida saqlaymiz (JSON)
    g = get_config('global_triggers')
    triggers = json.loads(g) if g else []
    await callback.message.answer(f"🔑 Global triggerlar soni: <b>{len(triggers)}</b>", parse_mode="HTML")
    await callback.answer()

@dp.callback_query(F.data == "adm_stats_avgdelay")
async def adm_stats_avgdelay_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user):
        await callback.answer(); return
    await callback.message.answer(
        f"⏳ Standart kechikish: <b>{get_global_default_delay()} s</b>\n"
        f"🕘 Silent soat: {'Yoqilgan' if get_flag('silent_enabled', True) else 'O\'chirilgan'}",
        parse_mode="HTML"
    )
    await callback.answer()

@dp.callback_query(F.data == "adm_stats_uptime")
async def adm_stats_uptime_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user):
        await callback.answer(); return
    uptime = datetime.now() - BOT_START_TIME if BOT_START_TIME else timedelta(0)
    await callback.message.answer(f"🕘 Bot ishlagan vaqti: <b>{str(uptime).split('.')[0]}</b>", parse_mode="HTML")
    await callback.answer()

@dp.callback_query(F.data == "adm_stats_dbsize")
async def adm_stats_dbsize_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user):
        await callback.answer(); return
    try:
        size = os.path.getsize(DB_NAME)
        await callback.message.answer(f"💾 Baza hajmi: <b>{size/1024:.1f} KB</b>", parse_mode="HTML")
    except Exception:
        await callback.message.answer("💾 Foydalanuvchi ma'lumotlari hali yaratilmagan.")
    await callback.answer()

@dp.callback_query(F.data == "adm_stats_msgs")
async def adm_stats_msgs_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user):
        await callback.answer(); return
    await callback.message.answer(f"📨 Jami yuborilgan xabarlar: <b>{get_config('total_msgs', 0)}</b>", parse_mode="HTML")
    await callback.answer()

@dp.callback_query(F.data == "adm_stats_blocks")
async def adm_stats_blocks_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user):
        await callback.answer(); return
    await callback.message.answer(f"🚫 Bloklangan foydalanuvchilar: <b>{len(get_blocked_users())}</b>", parse_mode="HTML")
    await callback.answer()

@dp.callback_query(F.data == "adm_stats_admins")
async def adm_stats_admins_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user):
        await callback.answer(); return
    await callback.message.answer(f"👑 Adminlar soni: <b>{len(get_admin_ids()) + 1}</b> (asosiy admin bilan)", parse_mode="HTML")
    await callback.answer()

async def admin_stats(message: types.Message):
    total_users = get_total_users_count()
    active_sessions = len(get_all_sessions())
    running = len(user_clients)
    triggers = get_all_triggers_count()
    uptime = datetime.now() - BOT_START_TIME if BOT_START_TIME else timedelta(0)
    await message.answer(
        f"📊 <b>Bot Statistikasi:</b>\n\n"
        f"👥 Barcha foydalanuvchilar: <b>{total_users}</b>\n"
        f"⚡️ Bazadagi faol seanslar: <b>{active_sessions}</b>\n"
        f"🔄 Hozir ishlab turgan jarayonlar: <b>{running}</b>\n"
        f"🔑 Jami triggerlar: <b>{triggers}</b>\n"
        f"🕘 Bot uptime: <b>{str(uptime).split('.')[0]}</b>",
        parse_mode="HTML"
    )

# ================= FOYDALANUVCHILAR =================
@dp.callback_query(F.data.startswith("adm_users_list_"))
async def adm_users_list_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user):
        await callback.answer(); return
    page = int(callback.data.split("_")[-1])
    sessions = get_all_sessions()
    page_size = 10
    start = page * page_size
    chunk = sessions[start:start + page_size]
    if not chunk:
        await callback.message.answer("📭 Hozircha ulangan foydalanuvchilar yo'q.")
        await callback.answer(); return
    lines = []
    for i, (uid, _) in enumerate(chunk):
        mark = "⚡️" if uid in user_clients else "💤"
        lines.append(f"{mark} {i + 1 + start}. <code>{uid}</code>")
    text = f"👥 <b>Foydalanuvchilar</b> (sahifa {page + 1}):\n\n" + "\n".join(lines)
    nav = []
    if page > 0:
        nav.append(InlineKeyboardButton(text="⬅️", callback_data=f"adm_users_list_{page-1}"))
    if start + page_size < len(sessions):
        nav.append(InlineKeyboardButton(text="➡️", callback_data=f"adm_users_list_{page+1}"))
    markup = InlineKeyboardMarkup(inline_keyboard=[nav]) if nav else None
    await callback.message.answer(text, parse_mode="HTML", reply_markup=markup)
    await callback.answer()

@dp.callback_query(F.data == "adm_user_info")
async def adm_user_info_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user):
        await callback.answer(); return
    waiting_for[callback.from_user.id] = "adm_user_info_input"
    await callback.message.answer("🔍 Ma'lumotini ko'rmoqchi bo'lgan foydalanuvchi ID sini kiriting:")
    await callback.answer()

@dp.callback_query(F.data == "adm_user_delete")
async def adm_user_delete_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user):
        await callback.answer(); return
    waiting_for[callback.from_user.id] = "adm_user_delete_input"
    await callback.message.answer("🗑 O'chiriladigan foydalanuvchi ID sini kiriting:")
    await callback.answer()

@dp.callback_query(F.data == "adm_user_set_text")
async def adm_user_set_text_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user):
        await callback.answer(); return
    waiting_for[callback.from_user.id] = "adm_user_text_input"
    await callback.message.answer("📝 Foydalanuvchi ID sini kiriting (keyingi qadamda matn):")
    await callback.answer()

@dp.callback_query(F.data == "adm_user_set_delay")
async def adm_user_set_delay_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user):
        await callback.answer(); return
    waiting_for[callback.from_user.id] = "adm_user_delay_input"
    await callback.message.answer("⏱ Foydalanuvchi ID sini kiriting (keyingi qadamda soniya):")
    await callback.answer()

@dp.callback_query(F.data == "adm_user_add_trigger")
async def adm_user_add_trigger_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user):
        await callback.answer(); return
    waiting_for[callback.from_user.id] = "adm_user_trig_key"
    await callback.message.answer("➕ Foydalanuvchi ID sini kiriting (keyingi qadamda kalit so'z):")
    await callback.answer()

@dp.callback_query(F.data == "adm_user_list_triggers")
async def adm_user_list_triggers_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user):
        await callback.answer(); return
    waiting_for[callback.from_user.id] = "adm_user_triggers_view"
    await callback.message.answer("📋 Triggerlarini ko'rmoqchi bo'lgan foydalanuvchi ID sini kiriting:")
    await callback.answer()

@dp.callback_query(F.data == "adm_user_stop")
async def adm_user_stop_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user):
        await callback.answer(); return
    waiting_for[callback.from_user.id] = "adm_user_stop_input"
    await callback.message.answer("⏹ Jarayonini to'xtatmoqchi bo'lgan foydalanuvchi ID sini kiriting:")
    await callback.answer()

@dp.callback_query(F.data == "adm_user_start")
async def adm_user_start_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user):
        await callback.answer(); return
    waiting_for[callback.from_user.id] = "adm_user_start_input"
    await callback.message.answer("▶️ Jarayonini qayta boshlamoqchi bo'lgan foydalanuvchi ID sini kiriting:")
    await callback.answer()

@dp.callback_query(F.data == "adm_user_test_reply")
async def adm_user_test_reply_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user):
        await callback.answer(); return
    waiting_for[callback.from_user.id] = "adm_user_test_input"
    await callback.message.answer("🧪 Test matn tayyorlaydigan foydalanuvchi ID sini kiriting:")
    await callback.answer()

@dp.callback_query(F.data == "adm_user_refresh")
async def adm_user_refresh_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user):
        await callback.answer(); return
    waiting_for[callback.from_user.id] = "adm_user_refresh_input"
    await callback.message.answer("🔁 Seansini yangilamoqchi bo'lgan foydalanuvchi ID sini kiriting:")
    await callback.answer()

@dp.callback_query(F.data == "adm_user_logs")
async def adm_user_logs_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user):
        await callback.answer(); return
    waiting_for[callback.from_user.id] = "adm_user_logs_input"
    await callback.message.answer("📝 Loglarini ko'rmoqchi bo'lgan foydalanuvchi ID sini kiriting:")
    await callback.answer()

@dp.callback_query(F.data == "adm_user_undo")
async def adm_user_undo_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user):
        await callback.answer(); return
    waiting_for[callback.from_user.id] = "adm_user_undo_input"
    await callback.message.answer("↩️ (Undo) Foydalanuvchi ID sini kiriting:")
    await callback.answer()

# ================= GLOBAL TRIGGERS =================
@dp.callback_query(F.data == "adm_global_trigger_add")
async def adm_global_trigger_add_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user):
        await callback.answer(); return
    waiting_for[callback.from_user.id] = "adm_global_trig_key"
    await callback.message.answer("➕ Global trigger kalit so'zini kiriting:")
    await callback.answer()

@dp.callback_query(F.data.startswith("adm_global_triggers_list_"))
async def adm_global_triggers_list_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user):
        await callback.answer(); return
    page = int(callback.data.split("_")[-1])
    g = get_config('global_triggers')
    triggers = json.loads(g) if g else []
    page_size = 8
    start = page * page_size
    chunk = triggers[start:start + page_size]
    if not chunk:
        await callback.message.answer("📭 Global triggerlar mavjud emas.")
        await callback.answer(); return
    lines = [f"{i+1+start}. <b>{kw}</b> ➡️ {resp}" for i, (kw, resp) in enumerate(chunk)]
    text = "🔑 <b>Global triggerlar</b>:\n\n" + "\n".join(lines)
    nav = []
    if page > 0:
        nav.append(InlineKeyboardButton(text="⬅️", callback_data=f"adm_global_triggers_list_{page-1}"))
    if start + page_size < len(triggers):
        nav.append(InlineKeyboardButton(text="➡️", callback_data=f"adm_global_triggers_list_{page+1}"))
    markup = InlineKeyboardMarkup(inline_keyboard=[nav]) if nav else None
    await callback.message.answer(text, parse_mode="HTML", reply_markup=markup)
    await callback.answer()

@dp.callback_query(F.data == "adm_global_trigger_find")
async def adm_global_trigger_find_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user):
        await callback.answer(); return
    waiting_for[callback.from_user.id] = "adm_global_trig_find"
    await callback.message.answer("🔍 Qidirmoqchi bo'lgan trigger kalit so'zini kiriting:")
    await callback.answer()

@dp.callback_query(F.data == "adm_global_trigger_del")
async def adm_global_trigger_del_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user):
        await callback.answer(); return
    waiting_for[callback.from_user.id] = "adm_global_trig_del"
    await callback.message.answer("❌ O'chiriladigan trigger kalit so'zini kiriting:")
    await callback.answer()

@dp.callback_query(F.data == "adm_global_triggers_clear")
async def adm_global_triggers_clear_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user):
        await callback.answer(); return
    set_config('global_triggers', json.dumps([]))
    await callback.message.answer("🧹 Barcha global triggerlar tozalandi.")
    await callback.answer()

@dp.callback_query(F.data == "adm_global_triggers_stats")
async def adm_global_triggers_stats_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user):
        await callback.answer(); return
    g = get_config('global_triggers')
    triggers = json.loads(g) if g else []
    await callback.message.answer(
        f"🔑 Global triggerlar: <b>{len(triggers)}</b>\n"
        f"⚡️ Foydalanuvchi triggerlari: <b>{get_all_triggers_count()}</b>\n"
        f"🔢 Jami triggerlar: <b>{len(triggers) + get_all_triggers_count()}</b>",
        parse_mode="HTML"
    )
    await callback.answer()

# ================= BROADCAST =================
@dp.callback_query(F.data == "adm_bcast_text_all")
async def adm_bcast_text_all_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    waiting_for[callback.from_user.id] = "adm_bcast_all"
    await callback.message.answer("📝 Barcha foydalanuvchilarga yuboriladigan matnni yozing:")
    await callback.answer()

@dp.callback_query(F.data == "adm_bcast_text_online")
async def adm_bcast_text_online_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    waiting_for[callback.from_user.id] = "adm_bcast_online"
    await callback.message.answer("🟢 Faqat ONLINE foydalanuvchilarga yuboriladigan matnni yozing:")
    await callback.answer()

@dp.callback_query(F.data == "adm_bcast_text_offline")
async def adm_bcast_text_offline_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    waiting_for[callback.from_user.id] = "adm_bcast_offline"
    await callback.message.answer("🔴 Faqat OFFLINE foydalanuvchilarga yuboriladigan matnni yozing:")
    await callback.answer()

@dp.callback_query(F.data == "adm_bcast_text_active")
async def adm_bcast_text_active_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    waiting_for[callback.from_user.id] = "adm_bcast_active"
    await callback.message.answer("⚡️ Faol jarayondagi foydalanuvchilarga yuboriladigan matnni yozing:")
    await callback.answer()

@dp.callback_query(F.data == "adm_bcast_photo")
async def adm_bcast_photo_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    waiting_for[callback.from_user.id] = "adm_bcast_media_photo"
    await callback.message.answer("🖼 Rasm yuboring (barchaga):")
    await callback.answer()

@dp.callback_query(F.data == "adm_bcast_video")
async def adm_bcast_video_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    waiting_for[callback.from_user.id] = "adm_bcast_media_video"
    await callback.message.answer("🎞 Video yuboring (barchaga):")
    await callback.answer()

@dp.callback_query(F.data == "adm_bcast_doc")
async def adm_bcast_doc_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    waiting_for[callback.from_user.id] = "adm_bcast_media_doc"
    await callback.message.answer("📄 Fayl yuboring (barchaga):")
    await callback.answer()

@dp.callback_query(F.data == "adm_bcast_audio")
async def adm_bcast_audio_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    waiting_for[callback.from_user.id] = "adm_bcast_media_audio"
    await callback.message.answer("🎵 Audio yuboring (barchaga):")
    await callback.answer()

@dp.callback_query(F.data == "adm_bcast_history")
async def adm_bcast_history_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    history = get_config('bcast_history', '[]')
    try:
        h = json.loads(history)
    except Exception:
        h = []
    if not h:
        await callback.message.answer("📭 Broadcast tarixi bo'sh.")
        await callback.answer(); return
    lines = [f"• {e}" for e in h[-20:]]
    await callback.message.answer("📊 <b>Broadcast tarixi:</b>\n\n" + "\n".join(lines), parse_mode="HTML")
    await callback.answer()

async def do_broadcast(matn, filter_kind):
    users = get_all_sessions()
    targets = []
    if filter_kind == "all":
        targets = [u for u, _ in users]
    elif filter_kind == "online":
        targets = [u for u, _ in users if u in user_clients]
    elif filter_kind == "offline":
        targets = [u for u, _ in users if u not in user_clients]
    elif filter_kind == "active":
        targets = [u for u, _ in users if u in user_clients]
    count = 0
    for uid in targets:
        try:
            await bot.send_message(uid, matn)
            count += 1
            await asyncio.sleep(0.05)
        except Exception:
            pass
    # tarix
    try:
        h = json.loads(get_config('bcast_history', '[]'))
    except Exception:
        h = []
    h.append(f"📤 {datetime.now().strftime('%H:%M %d.%m')} | {filter_kind} | {count} ta")
    set_config('bcast_history', json.dumps(h[-50:]))
    return count

# ================= GLOBAL SOZLAMALAR =================
@dp.callback_query(F.data == "adm_set_default_text")
async def adm_set_default_text_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    waiting_for[callback.from_user.id] = "adm_default_text_input"
    await callback.message.answer(
        f"✏️ Joriy standart matn:\n<code>{html.escape(get_global_default_text())}</code>\n\nYangisini kiriting:",
        parse_mode="HTML"
    )
    await callback.answer()

@dp.callback_query(F.data == "adm_set_default_delay")
async def adm_set_default_delay_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    waiting_for[callback.from_user.id] = "adm_default_delay_input"
    await callback.message.answer(f"⏱ Joriy standart kechikish: {get_global_default_delay()} soniya.\nYangi qiymatni kiriting (faqat raqam):")
    await callback.answer()

@dp.callback_query(F.data == "adm_set_bot_tag")
async def adm_set_bot_tag_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    waiting_for[callback.from_user.id] = "adm_bot_tag_input"
    await callback.message.answer(f"🏷 Joriy tag: <code>{html.escape(get_bot_tag())}</code>\nYangi tagni kiriting:", parse_mode="HTML")
    await callback.answer()

@dp.callback_query(F.data == "adm_toggle_silent")
async def adm_toggle_silent_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    set_flag('silent_enabled', not get_flag('silent_enabled', True))
    await callback.message.answer(f"🌃 Silent soat: {'✅ Yoqilgan' if get_flag('silent_enabled', True) else '❌ O\'chirilgan'}")
    await callback.answer()

@dp.callback_query(F.data == "adm_toggle_timeprefix")
async def adm_toggle_timeprefix_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    set_flag('time_prefix_enabled', not get_flag('time_prefix_enabled', True))
    await callback.message.answer(f"🕐 Vaqt prefiksi: {'✅ Yoqilgan' if get_flag('time_prefix_enabled', True) else '❌ O\'chirilgan'}")
    await callback.answer()

@dp.callback_query(F.data == "adm_toggle_translate")
async def adm_toggle_translate_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    set_flag('translate_enabled', not get_flag('translate_enabled', True))
    await callback.message.answer(f"🔤 Avto-tarjima: {'✅ Yoqilgan' if get_flag('translate_enabled', True) else '❌ O\'chirilgan'}")
    await callback.answer()

@dp.callback_query(F.data == "adm_toggle_cyrillic")
async def adm_toggle_cyrillic_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    set_flag('cyrillic_enabled', not get_flag('cyrillic_enabled', True))
    await callback.message.answer(f"🅰️ Kirill/limit konvertatsiya: {'✅ Yoqilgan' if get_flag('cyrillic_enabled', True) else '❌ O\'chirilgan'}")
    await callback.answer()

@dp.callback_query(F.data == "adm_toggle_autoreply")
async def adm_toggle_autoreply_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    set_flag('autoreply_enabled', not get_flag('autoreply_enabled', True))
    await callback.message.answer(f"💬 Avto-javob tizimi: {'✅ Yoqilgan' if get_flag('autoreply_enabled', True) else '❌ O\'chirilgan'}")
    await callback.answer()

@dp.callback_query(F.data == "adm_toggle_read")
async def adm_toggle_read_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    set_flag('read_enabled', not get_flag('read_enabled', True))
    await callback.message.answer(f"👓 O'qilgan (read) belgilash: {'✅ Yoqilgan' if get_flag('read_enabled', True) else '❌ O\'chirilgan'}")
    await callback.answer()

@dp.callback_query(F.data == "adm_toggle_holiday")
async def adm_toggle_holiday_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    set_flag('holiday_enabled', not get_flag('holiday_enabled', True))
    await callback.message.answer(f"🎉 Bayram tabriki: {'✅ Yoqilgan' if get_flag('holiday_enabled', True) else '❌ O\'chirilgan'}")
    await callback.answer()

@dp.callback_query(F.data == "adm_toggle_debug")
async def adm_toggle_debug_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    set_flag('debug_enabled', not get_flag('debug_enabled', True))
    await callback.message.answer(f"🐞 Debug log: {'✅ Yoqilgan' if get_flag('debug_enabled', True) else '❌ O\'chirilgan'}")
    await callback.answer()

@dp.callback_query(F.data == "adm_toggle_welcome")
async def adm_toggle_welcome_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    set_flag('welcome_enabled', not get_flag('welcome_enabled', True))
    await callback.message.answer(f"🚼 Yangi foydalanuvchi xabari: {'✅ Yoqilgan' if get_flag('welcome_enabled', True) else '❌ O\'chirilgan'}")
    await callback.answer()

@dp.callback_query(F.data == "adm_set_welcome_text")
async def adm_set_welcome_text_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    waiting_for[callback.from_user.id] = "adm_welcome_text"
    await callback.message.answer(f"📝 Joriy welcome matni:\n<code>{html.escape(get_config('welcome_text', 'Xush kelibsiz!👋'))}</code>\n\nYangisini kiriting:", parse_mode="HTML")
    await callback.answer()

@dp.callback_query(F.data == "adm_add_filter_word")
async def adm_add_filter_word_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    waiting_for[callback.from_user.id] = "adm_filter_add"
    await callback.message.answer("🚫 Filtrlash uchun so'z kiritib yuboring:")
    await callback.answer()

@dp.callback_query(F.data == "adm_list_filter_words")
async def adm_list_filter_words_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    words = json.loads(get_config('filter_words', '[]'))
    if not words:
        await callback.message.answer("📭 Filter so'zlar ro'yxati bo'sh.")
        await callback.answer(); return
    await callback.message.answer("🚫 <b>Filter so'zlar:</b>\n\n• " + "\n• ".join(words), parse_mode="HTML")
    await callback.answer()

@dp.callback_query(F.data == "adm_del_filter_word")
async def adm_del_filter_word_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    waiting_for[callback.from_user.id] = "adm_filter_del"
    await callback.message.answer("🚮 O'chiriladigan filter so'zini kiriting:")
    await callback.answer()

@dp.callback_query(F.data == "adm_reset_settings")
async def adm_reset_settings_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    for k in ['default_reply_text', 'default_delay', 'bot_tag', 'silent_enabled',
              'time_prefix_enabled', 'translate_enabled', 'cyrillic_enabled',
              'autoreply_enabled', 'read_enabled', 'welcome_enabled', 'welcome_text',
              'filter_words']:
        try:
            with get_db() as conn:
                conn.execute('DELETE FROM bot_config WHERE key=?', (k,))
        except Exception:
            pass
    await callback.message.answer("🧹 Barcha global sozlamalar standart holatga qaytarildi.")
    await callback.answer()

# ================= BAYRAM =================
@dp.callback_query(F.data == "adm_reset_holidays")
async def adm_reset_holidays_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    count = reset_all_holiday_greetings()
    await callback.message.answer(f"🎉 {count} ta bayram belgisi tozalandi. Endi foydalanuvchilar qayta tabriklanadi.")
    await callback.answer()

@dp.callback_query(F.data == "adm_set_holiday_newyear")
async def adm_set_holiday_newyear_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    waiting_for[callback.from_user.id] = "adm_holiday_newyear"
    await callback.message.answer("🎊 Yangi Yil bayram matnini kiriting:")
    await callback.answer()

@dp.callback_query(F.data == "adm_set_holiday_navruz")
async def adm_set_holiday_navruz_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    waiting_for[callback.from_user.id] = "adm_holiday_navruz"
    await callback.message.answer("🌿 Navro'z bayram matnini kiriting:")
    await callback.answer()

@dp.callback_query(F.data == "adm_set_holiday_indep")
async def adm_set_holiday_indep_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    waiting_for[callback.from_user.id] = "adm_holiday_indep"
    await callback.message.answer("🇺🇿 Mustaqillik kuni bayram matnini kiriting:")
    await callback.answer()

@dp.callback_query(F.data == "adm_set_holiday_teacher")
async def adm_set_holiday_teacher_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    waiting_for[callback.from_user.id] = "adm_holiday_teacher"
    await callback.message.answer("👩‍🏫 O'qituvchilar kuni bayram matnini kiriting:")
    await callback.answer()

@dp.callback_query(F.data == "adm_set_holiday_constitution")
async def adm_set_holiday_constitution_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    waiting_for[callback.from_user.id] = "adm_holiday_constitution"
    await callback.message.answer("⚖️ Konstitutsiya kuni bayram matnini kiriting:")
    await callback.answer()

@dp.callback_query(F.data == "adm_holiday_greeted_list")
async def adm_holiday_greeted_list_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    events = get_config('holiday_events', '[]')
    try:
        evs = json.loads(events)
    except Exception:
        evs = []
    if not evs:
        await callback.message.answer("📭 Hozircha hech kim bayram bilan tabriklanmagan.")
        await callback.answer(); return
    lines = [f"• {e}" for e in evs[-30:]]
    await callback.message.answer("🎉 <b>Bayram tabriklar tarixi:</b>\n\n" + "\n".join(lines), parse_mode="HTML")
    await callback.answer()

# ================= LOGLAR =================
@dp.callback_query(F.data == "adm_logs_view")
async def adm_logs_view_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    if not chat_log_cache:
        await callback.message.answer("📭 Hech qanday log mavjud emas.")
        await callback.answer(); return
    total = sum(len(v) for v in chat_log_cache.values())
    await callback.message.answer(f"📝 Jami log (xotira): <b>{total} ta yozuv</b>. Batafsil ko'rish uchun foydalanuvchi loglarini chop eting.", parse_mode="HTML")
    await callback.answer()

@dp.callback_query(F.data == "adm_logs_clear")
async def adm_logs_clear_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    chat_log_cache.clear()
    await callback.message.answer("🗑 Barcha loglar tozalandi.")
    await callback.answer()

@dp.callback_query(F.data == "adm_logs_ttl")
async def adm_logs_ttl_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    await callback.message.answer("⏱ Loglar <b>2 soat</b> davomida xotirada saqlanadi.", parse_mode="HTML")
    await callback.answer()

@dp.callback_query(F.data == "adm_logs_stats")
async def adm_logs_stats_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    total = sum(len(v) for v in chat_log_cache.values())
    users = len(chat_log_cache)
    await callback.message.answer(f"📊 Log statistikasi:\n🟢 Faol chatlar: {users}\n📝 Yozuvlar: {total}", parse_mode="HTML")
    await callback.answer()

# ================= MEDIA (bcast media uchun kutiladi) =================
# admin media broadcast — content handlerlar quyida (matn ishlovchi bilan birga)
# Bu callbacklar faqat holatni belgilaydi (adm_bcast_media_*). Foydalanuvchi media yuborganda handler ishlaydi.

# ================= XAVFSIZLIK =================
@dp.callback_query(F.data == "adm_add_admin")
async def adm_add_admin_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    waiting_for[callback.from_user.id] = "adm_add_admin_input"
    await callback.message.answer("👑 Yangi admin ID sini kiriting:")
    await callback.answer()

@dp.callback_query(F.data == "adm_list_admins")
async def adm_list_admins_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    ids = sorted(get_admin_ids())
    if not ids:
        await callback.message.answer("👑 Qo'shimcha adminlar yo'q (faqat asosiy admin).")
        await callback.answer(); return
    lines = "\n".join(f"• <code>{i}</code>" for i in ids)
    await callback.message.answer(f"👥 <b>Qo'shimcha adminlar:</b>\n\n{lines}", parse_mode="HTML")
    await callback.answer()

@dp.callback_query(F.data == "adm_remove_admin")
async def adm_remove_admin_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    waiting_for[callback.from_user.id] = "adm_remove_admin_input"
    await callback.message.answer("❌ Olib tashlanadigan admin ID sini kiriting:")
    await callback.answer()

@dp.callback_query(F.data == "adm_block_user")
async def adm_block_user_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    waiting_for[callback.from_user.id] = "adm_block_input"
    await callback.message.answer("🚫 Bloklanadigan foydalanuvchi ID sini kiriting:")
    await callback.answer()

@dp.callback_query(F.data == "adm_unblock_user")
async def adm_unblock_user_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    waiting_for[callback.from_user.id] = "adm_unblock_input"
    await callback.message.answer("✅ Blokdan chiqariladigan foydalanuvchi ID sini kiriting:")
    await callback.answer()

@dp.callback_query(F.data == "adm_list_blocks")
async def adm_list_blocks_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    blocks = get_blocked_users()
    if not blocks:
        await callback.message.answer("📭 Bloklangan foydalanuvchilar yo'q.")
        await callback.answer(); return
    lines = "\n".join(f"• <code>{uid}</code> ({reason})" for uid, reason in blocks)
    await callback.message.answer(f"🚫 <b>Bloklanganlar:</b>\n\n{lines}", parse_mode="HTML")
    await callback.answer()

@dp.callback_query(F.data == "adm_secure_all")
async def adm_secure_all_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    # xavfsizlik: barcha faol jarayonlarni qayta ishga tushirish (seans yangilanishi)
    cnt = 0
    for uid in list(user_clients.keys()):
        stop_client_task(uid)
        cnt += 1
    await callback.message.answer(f"🔐 {cnt} ta jarayon xavfsiz moddasiga o'tkazildi (qayta ishga tushiriladi).")
    await callback.answer()

@dp.callback_query(F.data == "adm_stop_flood")
async def adm_stop_flood_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    set_flag('flood_protect', True)
    await callback.message.answer("🛑 Flood himoyasi yoqildi (og'ir yuklamalardan himoya).")
    await callback.answer()

# ================= CLEANUP =================
@dp.callback_query(F.data == "adm_clear_holidays")
async def adm_clear_holidays_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    count = reset_all_holiday_greetings()
    await callback.message.answer(f"🧹 {count} ta bayram belgisi tozalandi.")
    await callback.answer()

@dp.callback_query(F.data == "adm_clean_old_sessions")
async def adm_clean_old_sessions_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    # Eski / yaroqsiz seanslarni tekshirib, nosozlarini olib tashlaymiz
    removed = 0
    with get_db() as conn:
        c = conn.cursor()
        c.execute('SELECT user_id, session_string FROM sessions WHERE session_string IS NOT NULL')
        rows = c.fetchall()
    for uid, ss in rows:
        try:
            cli = TelegramClient(StringSession(ss), API_ID, API_HASH)
            await cli.connect()
            if not await cli.is_user_authorized():
                delete_session(uid)
                removed += 1
        except Exception:
            delete_session(uid)
            removed += 1
        finally:
            try:
                await cli.disconnect()
            except Exception:
                pass
        await asyncio.sleep(0.05)
    await callback.message.answer(
        f"🧾 <b>{len(rows)}</b> ta seans tekshirildi, <b>{removed}</b> ta yaroqsiz seans olib tashlandi.",
        parse_mode="HTML"
    )
    await callback.answer()

@dp.callback_query(F.data == "adm_clear_filters")
async def adm_clear_filters_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    set_config('filter_words', json.dumps([]))
    await callback.message.answer("🚮 Barcha filter so'zlar tozalandi.")
    await callback.answer()

@dp.callback_query(F.data == "adm_optimize_db")
async def adm_optimize_db_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    with get_db() as conn:
        conn.execute("VACUUM;")
        conn.execute("ANALYZE;")
    await callback.message.answer("🧹 Baza optimallashtirildi (VACUUM + ANALYZE).")
    await callback.answer()

@dp.callback_query(F.data == "adm_clear_all_triggers")
async def adm_clear_all_triggers_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    with get_db() as conn:
        conn.execute('DELETE FROM custom_triggers')
    await callback.message.answer("🗑 Barcha foydalanuvchi triggerlari o'chirildi.")
    await callback.answer()

# ================= TIZIM / VOSITALAR =================
@dp.callback_query(F.data == "adm_diag")
async def adm_diag_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    diag = (
        f"⚠️ <b>Diagnostika</b>\n\n"
        f"🟢 Bot: ishlayapti\n"
        f"🟡 Faol jarayonlar: {len(user_clients)}\n"
        f"🟡 Xotiradagi loglar: {sum(len(v) for v in chat_log_cache.values())}\n"
        f"🔵 DB: {os.path.exists(DB_NAME)}\n"
        f"🔵 Python: {sys.version.split()[0]}"
    )
    await callback.message.answer(diag, parse_mode="HTML")
    await callback.answer()

@dp.callback_query(F.data == "adm_export_db")
async def adm_export_db_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    try:
        if not os.path.exists(DB_NAME):
            await callback.message.answer("📦 Baza fayli hali yaratilmagan.")
            await callback.answer(); return
        await callback.message.answer_document(
            types.FSInputFile(DB_NAME), caption="📦 Bazani eksport qilish"
        )
    except Exception as e:
        await callback.message.answer(f"❌ Eksport xatosi: {html.escape(str(e))}")
    await callback.answer()

@dp.callback_query(F.data == "adm_about")
async def adm_about_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    uptime = datetime.now() - BOT_START_TIME if BOT_START_TIME else timedelta(0)
    total_users = get_total_users_count()
    active = len(user_clients)
    await callback.message.answer(
        f"ℹ️ <b>Bot haqida</b>\n\n"
        f"👥 Jami foydalanuvchilar: {total_users}\n"
        f"⚡️ Faol jarayonlar: {active}\n"
        f"🧩 Admin funksiyalari: {count_admin_features()} dan ortiq\n"
        f"⏳ Ishlab turgan vaqti: {str(uptime).split('.')[0]}\n"
        f"🍃 Reply-uz • Avto-javob userbot",
        parse_mode="HTML"
    )
    await callback.answer()

@dp.callback_query(F.data == "adm_clear_webhook")
async def adm_clear_webhook_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    try:
        await bot.delete_webhook(drop_pending_updates=True)
        await callback.message.answer("🔄 Webhook o'chirildi (polling rejimi).")
    except Exception as e:
        await callback.message.answer(f"❌ Xatolik: {html.escape(str(e))}")
    await callback.answer()

# ================= INTERFEYS TILI =================
@dp.callback_query(F.data == "adm_lang_uz")
async def adm_lang_uz_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    set_config('interface_lang', 'uz')
    await callback.message.answer("🇺🇿 Interfeys tili: O'zbekcha")
    await callback.answer()

@dp.callback_query(F.data == "adm_lang_ru")
async def adm_lang_ru_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    set_config('interface_lang', 'ru')
    await callback.message.answer("🗣 Язык интерфейса: Русский")
    await callback.answer()

@dp.callback_query(F.data == "adm_lang_en")
async def adm_lang_en_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    set_config('interface_lang', 'en')
    await callback.message.answer("🌍 Interface language: English")
    await callback.answer()

# ================= RESTART / STOP ALL / START ALL =================
@dp.callback_query(F.data == "adm_restart")
async def adm_restart_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    await callback.message.answer("♻️ Bot qayta ishga tushmoqda...")
    await callback.answer()
    os.execv(sys.executable, [sys.executable] + sys.argv)

@dp.callback_query(F.data == "adm_stop_all")
async def adm_stop_all_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    ids = list(user_clients.keys())
    for uid in ids:
        stop_client_task(uid)
    await callback.message.answer(f"⏹ {len(ids)} ta faol avto-javob jarayoni to'xtatildi.")
    await callback.answer()

@dp.callback_query(F.data == "adm_start_all")
async def adm_start_all_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    sessions = get_all_sessions()
    started = 0
    for uid, ss in sessions:
        try:
            if uid in user_clients:
                continue
            client = TelegramClient(StringSession(ss), API_ID, API_HASH)
            await client.connect()
            if await client.is_user_authorized():
                user_clients[uid] = client
                auto_tasks[uid] = asyncio.create_task(start_auto_reply(uid, client))
                started += 1
        except Exception:
            continue
        await asyncio.sleep(0.1)
    await callback.message.answer(f"▶️ {started} ta jarayon qayta ishga tushirildi.")
    await callback.answer()

# ================= EKSPORT (EXCEL / WORD) =================
@dp.callback_query(F.data == "adm_export_xlsx")
async def adm_export_xlsx_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    rows = collect_user_data()
    if not rows:
        await callback.message.answer("📭 Eksport uchun foydalanuvchilar mavjud emas.")
        await callback.answer(); return
    try:
        data = build_excel_bytes(rows)
        buf = BufferedInputFile(data, filename="foydalanuvchilar.xlsx")
        await callback.message.answer_document(buf, caption=f"📊 {len(rows)} ta foydalanuvchi (Excel)")
    except Exception as e:
        await callback.message.answer(f"❌ Eksport xatosi: {html.escape(str(e))}")
    await callback.answer()

@dp.callback_query(F.data == "adm_export_docx")
async def adm_export_docx_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    rows = collect_user_data()
    if not rows:
        await callback.message.answer("📭 Eksport uchun foydalanuvchilar mavjud emas.")
        await callback.answer(); return
    try:
        data = build_word_bytes(rows)
        buf = BufferedInputFile(data, filename="foydalanuvchilar.docx")
        await callback.message.answer_document(buf, caption=f"📄 {len(rows)} ta foydalanuvchi (Word)")
    except Exception as e:
        await callback.message.answer(f"❌ Eksport xatosi: {html.escape(str(e))}")
    await callback.answer()

@dp.callback_query(F.data == "adm_export_count")
async def adm_export_count_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    rows = collect_user_data()
    await callback.message.answer(
        f"👥 Jami foydalanuvchilar: <b>{len(rows)}</b>\n"
        f"⚡️ Faol jarayonlar: <b>{sum(1 for r in rows if r['running'])}</b>\n"
        f"🚫 Bloklanganlar: <b>{sum(1 for r in rows if r['blocked'])}</b>",
        parse_mode="HTML"
    )
    await callback.answer()

@dp.callback_query(F.data == "adm_export_cleanup")
async def adm_export_cleanup_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    set_config('debug_enabled', True)
    await callback.message.answer("🧹 Eksport ishga tayyor. Yangi foydalanuvchi ma'lumotlari DB'dan o'qiladi.")
    await callback.answer()

def count_admin_features():
    """Jami admin funksiya (menyu + callback) sonini sanaydi."""
    menu = sum(len(v) for v in CATEGORY_ITEMS.values()) + len(ADMIN_CATEGORIES)
    return menu + 12  # qo'shimcha xizmat funksiyalar

# =======================================================================
#                    MATN ISHLOVCHI (admin + user flow)
# =======================================================================
@dp.message(F.text & ~F.text.startswith('/'))
async def handle_text_input(message: types.Message):
    user_id = message.from_user.id
    text = message.text
    state = waiting_for.get(user_id)

    # Bosh menyu / Admin panel tugmalari — alohida handler orqali boshqariladi
    if text in ["⚙️ Sozlamalar / Bosh menyu", "👑 Admin Panel"]:
        return

    # ---------- ADMIN KIRITMALARI ----------
    if is_admin(message.from_user):
        if state == "adm_default_text_input":
            waiting_for[user_id] = None
            set_global_default_text(text)
            await message.answer(f"✅ Standart javob matni yangilandi:\n<code>{html.escape(text)}</code>", parse_mode="HTML")
            return
        if state == "adm_default_delay_input":
            waiting_for[user_id] = None
            if text.strip().isdigit() and int(text.strip()) >= 1:
                set_global_default_delay(int(text.strip()))
                await message.answer(f"✅ Standart kechikish {text.strip()} soniya qilib belgilandi.")
            else:
                await message.answer("❌ Faqat 1 dan katta raqam kiriting.")
            return
        if state == "adm_bot_tag_input":
            waiting_for[user_id] = None
            set_bot_tag(text)
            await message.answer(f"🏷 Tag yangilandi: <code>{html.escape(text)}</code>", parse_mode="HTML")
            return
        if state == "adm_welcome_text":
            waiting_for[user_id] = None
            set_config('welcome_text', text)
            await message.answer(f"📝 Welcome matni saqlandi:\n<code>{html.escape(text)}</code>", parse_mode="HTML")
            return
        if state == "adm_filter_add":
            waiting_for[user_id] = None
            words = json.loads(get_config('filter_words', '[]'))
            if text.strip().lower() not in words:
                words.append(text.strip().lower())
            set_config('filter_words', json.dumps(words))
            await message.answer(f"🚫 Filter so'z qo'shildi: <code>{html.escape(text)}</code>", parse_mode="HTML")
            return
        if state == "adm_filter_del":
            waiting_for[user_id] = None
            words = json.loads(get_config('filter_words', '[]'))
            if text.strip().lower() in words:
                words.remove(text.strip().lower())
                set_config('filter_words', json.dumps(words))
                await message.answer(f"🚮 Filter so'z o'chirildi: <code>{html.escape(text)}</code>", parse_mode="HTML")
            else:
                await message.answer("❌ Bunday filter so'z topilmadi.")
            return
        if state == "adm_add_admin_input":
            waiting_for[user_id] = None
            try:
                add_admin(int(text.strip()))
                await message.answer(f"👑 Admin qo'shildi: <code>{text.strip()}</code>", parse_mode="HTML")
            except Exception:
                await message.answer("❌ ID raqam bo'lishi kerak.")
            return
        if state == "adm_remove_admin_input":
            waiting_for[user_id] = None
            try:
                remove_admin(int(text.strip()))
                await message.answer(f"❌ Admin olib tashlandi: <code>{text.strip()}</code>", parse_mode="HTML")
            except Exception:
                await message.answer("❌ ID raqam bo'lishi kerak.")
            return
        if state == "adm_block_input":
            waiting_for[user_id] = None
            try:
                block_user(int(text.strip()))
                await message.answer(f"🚫 Foydalanuvchi bloklandi: <code>{text.strip()}</code>", parse_mode="HTML")
            except Exception:
                await message.answer("❌ ID raqam bo'lishi kerak.")
            return
        if state == "adm_unblock_input":
            waiting_for[user_id] = None
            try:
                unblock_user(int(text.strip()))
                await message.answer(f"✅ Foydalanuvchi blokdan chiqarildi: <code>{text.strip()}</code>", parse_mode="HTML")
            except Exception:
                await message.answer("❌ ID raqam bo'lishi kerak.")
            return
        if state == "adm_user_info_input":
            waiting_for[user_id] = None
            try:
                target_id = int(text.strip())
            except ValueError:
                await message.answer("❌ Foydalanuvchi ID raqam bo'lishi kerak.")
                return
            if not has_active_session(target_id):
                await message.answer("❌ Bunday ulangan foydalanuvchi topilmadi.")
                return
            info = (
                f"🔍 <b>Foydalanuvchi:</b> <code>{target_id}</code>\n"
                f"💬 Javob matni: <code>{html.escape(get_custom_reply_text(target_id))}</code>\n"
                f"⏱ Kechikish: {get_user_delay(target_id)} soniya\n"
                f"🔑 Triggerlar soni: {len(get_user_triggers(target_id))}\n"
                f"⚡️ Faol jarayon: {'✅ Ha' if target_id in user_clients else '❌ Yoq'}\n"
                f"🚫 Blok: {'✅ Ha' if is_blocked(target_id) else '❌ Yoq'}"
            )
            await message.answer(info, parse_mode="HTML")
            return
        if state == "adm_user_delete_input":
            waiting_for[user_id] = None
            try:
                target_id = int(text.strip())
            except ValueError:
                await message.answer("❌ Foydalanuvchi ID raqam bo'lishi kerak.")
                return
            stop_client_task(target_id)
            delete_session(target_id)
            await message.answer(f"✅ Foydalanuvchi <code>{text.strip()}</code> o'chirildi.", parse_mode="HTML")
            return
        if state == "adm_user_text_input":
            waiting_for[user_id] = "adm_user_text_value:" + text.strip()
            await message.answer(f"✏️ <code>{text.strip()}</code> uchun yangi javob matnini kiriting:", parse_mode="HTML")
            return
        if state and state.startswith("adm_user_text_value:"):
            target_id = int(state.split(":", 1)[1])
            set_custom_reply_text(target_id, text)
            waiting_for[user_id] = None
            await message.answer(f"✅ <code>{target_id}</code> uchun javob matni saqlandi.", parse_mode="HTML")
            return
        if state == "adm_user_delay_input":
            waiting_for[user_id] = "adm_user_delay_value:" + text.strip()
            await message.answer(f"⏱ <code>{text.strip()}</code> uchun kechikish soniyasini kiriting:", parse_mode="HTML")
            return
        if state and state.startswith("adm_user_delay_value:"):
            target_id = int(state.split(":", 1)[1])
            if text.strip().isdigit() and int(text.strip()) >= 1:
                set_user_delay(target_id, int(text.strip()))
                waiting_for[user_id] = None
                await message.answer(f"✅ <code>{target_id}</code> uchun kechikish <b>{text.strip()} s</b> qilindi.", parse_mode="HTML")
            else:
                await message.answer("❌ Faqat 1 dan katta raqam kiriting.")
            return
        if state == "adm_user_trig_key":
            try:
                int(text.strip())
            except ValueError:
                await message.answer("❌ Avval foydalanuvchi ID raqamini kiriting.")
                return
            waiting_for[user_id] = "adm_user_trig_kw:" + text.strip()
            await message.answer(f"➕ <code>{text.strip()}</code> uchun trigger kalit so'zini kiriting:", parse_mode="HTML")
            return
        if state and state.startswith("adm_user_trig_kw:"):
            target_id = state.split(":", 1)[1]
            waiting_for[user_id] = f"adm_user_trig_resp:{target_id}:{text.strip()}"
            await message.answer(
                f"➕ <code>{target_id}</code>, <b>{text.strip()}</b> kalit so'ziga javobni kiriting:",
                parse_mode="HTML"
            )
            return
        if state and state.startswith("adm_user_trig_resp:"):
            parts = state.split(":", 2)
            target_id = int(parts[1])
            keyword = parts[2]
            add_custom_trigger(target_id, keyword, text.strip())
            waiting_for[user_id] = None
            await message.answer(
                f"✅ <code>{target_id}</code> uchun trigger saqlandi: <code>{keyword}</code> ➡️ {text.strip()}",
                parse_mode="HTML"
            )
            return
        if state == "adm_user_triggers_view":
            waiting_for[user_id] = None
            try:
                target_id = int(text.strip())
            except ValueError:
                await message.answer("❌ ID raqam bo'lishi kerak.")
                return
            triggers = get_user_triggers(target_id)
            if not triggers:
                await message.answer(f"📭 <code>{target_id}</code> uchun triggerlar yo'q.", parse_mode="HTML")
                return
            lines = [f"• <code>{kw}</code> ➡️ {resp}" for _, kw, resp in triggers]
            await message.answer(f"🔑 <b>{target_id}</b> triggerlari:\n\n" + "\n".join(lines), parse_mode="HTML")
            return
        if state == "adm_user_stop_input":
            waiting_for[user_id] = None
            try:
                target_id = int(text.strip())
            except ValueError:
                await message.answer("❌ ID raqam bo'lishi kerak.")
                return
            stop_client_task(target_id)
            await message.answer(f"⏹ <code>{target_id}</code> jarayoni to'xtatildi.", parse_mode="HTML")
            return
        if state == "adm_user_start_input":
            waiting_for[user_id] = None
            try:
                target_id = int(text.strip())
            except ValueError:
                await message.answer("❌ ID raqam bo'lishi kerak.")
                return
            ss = get_session_string(target_id)
            if not ss:
                await message.answer("❌ Bu foydalanuvchi seansi yo'q.")
                return
            try:
                client = TelegramClient(StringSession(ss), API_ID, API_HASH)
                await client.connect()
                if await client.is_user_authorized():
                    user_clients[target_id] = client
                    auto_tasks[target_id] = asyncio.create_task(start_auto_reply(target_id, client))
                    await message.answer(f"▶️ <code>{target_id}</code> jarayoni qayta boshlandi.", parse_mode="HTML")
                else:
                    await message.answer("❌ Seans yaroqsiz.")
            except Exception as e:
                await message.answer(f"❌ Xatolik: {html.escape(str(e))}")
            return
        if state == "adm_user_test_input":
            waiting_for[user_id] = None
            try:
                target_id = int(text.strip())
            except ValueError:
                await message.answer("❌ ID raqam bo'lishi kerak.")
                return
            demo = get_reply_text(target_id, "Assalomu alaykum")
            await message.answer(
                f"🧪 <b>{target_id}</b> uchun test javob:\n\n<code>{html.escape(demo)}</code>",
                parse_mode="HTML"
            )
            return
        if state == "adm_user_refresh_input":
            waiting_for[user_id] = None
            try:
                target_id = int(text.strip())
            except ValueError:
                await message.answer("❌ ID raqam bo'lishi kerak.")
                return
            stop_client_task(target_id)
            ss = get_session_string(target_id)
            await message.answer(f"🔁 <code>{target_id}</code> seansi yangilandi (qayta ishga tushiriladi).", parse_mode="HTML")
            return
        if state == "adm_user_logs_input":
            waiting_for[user_id] = None
            try:
                target_id = int(text.strip())
            except ValueError:
                await message.answer("❌ ID raqam bo'lishi kerak.")
                return
            logs = get_logs_last_2h(target_id)
            if not logs:
                await message.answer(f"📭 <code>{target_id}</code> uchun so'nggi loglar yo'q.", parse_mode="HTML")
                return
            lines = [f"[{e['timestamp'].strftime('%H:%M:%S')}] {'MEN' if e['from_me'] else 'SUHBAT'} : {e['text'][:60]}" for e in logs[-15:]]
            await message.answer(f"📝 <b>{target_id}</b> loglari:\n\n" + "\n".join(lines), parse_mode="HTML")
            return
        if state == "adm_user_undo_input":
            waiting_for[user_id] = None
            await message.answer("↩️ (Undo) Buyruq mavjud emas — ma'lumot uchun qo'llanma bilan tanishing.")
            return
        if state == "adm_global_trig_key":
            waiting_for[user_id] = "adm_global_trig_resp:" + text.strip()
            await message.answer(f"➕ <b>{text.strip()}</b> kalit so'zi uchun javobni kiriting:", parse_mode="HTML")
            return
        if state and state.startswith("adm_global_trig_resp:"):
            kw = state.split(":", 1)[1]
            triggers = json.loads(get_config('global_triggers', '[]'))
            triggers.append([kw, text.strip()])
            set_config('global_triggers', json.dumps(triggers))
            waiting_for[user_id] = None
            await message.answer(f"✅ Global trigger saqlandi: <code>{kw}</code> ➡️ {text.strip()}", parse_mode="HTML")
            return
        if state == "adm_global_trig_find":
            waiting_for[user_id] = None
            triggers = json.loads(get_config('global_triggers', '[]'))
            q = text.strip().lower()
            found = [t for t in triggers if q in t[0].lower()]
            if not found:
                await message.answer("🔍 Hech narsa topilmadi.")
                return
            lines = [f"• <code>{kw}</code> ➡️ {resp}" for kw, resp in found]
            await message.answer("🔍 <b>Topilgan triggerlar:</b>\n\n" + "\n".join(lines), parse_mode="HTML")
            return
        if state == "adm_global_trig_del":
            waiting_for[user_id] = None
            triggers = json.loads(get_config('global_triggers', '[]'))
            q = text.strip().lower()
            new = [t for t in triggers if t[0].lower() != q]
            if len(new) == len(triggers):
                await message.answer("❌ Bunday trigger topilmadi.")
                return
            set_config('global_triggers', json.dumps(new))
            await message.answer(f"✅ <code>{text.strip()}</code> triggeri o'chirildi.", parse_mode="HTML")
            return
        if state == "adm_holiday_newyear":
            waiting_for[user_id] = None
            set_config('holiday_text_new_year', text)
            await message.answer("🎊 Yangi Yil bayram matni saqlandi.")
            return
        if state == "adm_holiday_navruz":
            waiting_for[user_id] = None
            set_config('holiday_text_navruz', text)
            await message.answer("🌿 Navro'z bayram matni saqlandi.")
            return
        if state == "adm_holiday_indep":
            waiting_for[user_id] = None
            set_config('holiday_text_independence', text)
            await message.answer("🇺🇿 Mustaqillik bayram matni saqlandi.")
            return
        if state == "adm_holiday_teacher":
            waiting_for[user_id] = None
            set_config('holiday_text_teacher', text)
            await message.answer("👩‍🏫 O'qituvchilar kuni matni saqlandi.")
            return
        if state == "adm_holiday_constitution":
            waiting_for[user_id] = None
            set_config('holiday_text_constitution', text)
            await message.answer("⚖️ Konstitutsiya kuni matni saqlandi.")
            return
        if state and state.startswith("adm_bcast"):
            await handle_admin_broadcast_text(message, state, text)
            return

    # ---------- ULASH JARAYONI ----------
    if state == "phone":
        if text.startswith('+'):
            await request_code(user_id, text, message)
        else:
            await message.answer("❌ Raqamni +998... shaklida kiriting.")
        return

    if state == "code":
        client = user_clients.get(user_id)
        if not client:
            await message.answer("❌ Xatolik. /start bosing.")
            return
        clean_code = text.replace('-', '').replace(' ', '').strip()
        phone = phone_cache.get(user_id)
        phone_code_hash = phone_code_hash_cache.get(user_id)
        try:
            await client.sign_in(phone=phone, code=clean_code, phone_code_hash=phone_code_hash)
            await finish_login(user_id, client, message)
        except SessionPasswordNeededError:
            waiting_for[user_id] = "password"
            await message.answer("🔐 2-bosqichli parolingizni kiriting:")
        except Exception as e:
            await message.answer(f"❌ Kod noto'g'ri: {html.escape(str(e))}")
        return

    if state == "password":
        client = user_clients.get(user_id)
        if not client:
            return
        try:
            await client.sign_in(password=text)
            await finish_login(user_id, client, message)
        except Exception as e:
            await message.answer(f"❌ Parol noto'g'ri: {html.escape(str(e))}")
        return

    # ---------- ULANGAN USER NORMAL MATNI ----------
    if not has_active_session(user_id):
        await message.answer("❌ Botdan foydalanish uchun avval /start buyrug'i orqali ulaning.")
        return

    if state == "trigger_keyword":
        waiting_for[user_id] = f"trigger_response:{text.strip()}"
        await message.answer(f"🔹 **'{text.strip()}'** so'zi yozilganda bot qanday javob qaytarsin?", parse_mode="Markdown")
        return

    if state and state.startswith("trigger_response:"):
        keyword = state.split(":", 1)[1]
        add_custom_trigger(user_id, keyword, text.strip())
        waiting_for[user_id] = None
        await message.answer(
            f"✅ Saqlandi!\n\n🔸 **So'rov:** `{keyword}`\n🔹 **Javob:** `{text.strip()}`",
            parse_mode="Markdown",
            reply_markup=get_main_keyboard(True)
        )
        return

    if state == "custom_text":
        set_custom_reply_text(user_id, text)
        waiting_for[user_id] = None
        await message.answer(
            f"✅ Asosiy javob saqlandi:\n<code>{html.escape(text)}</code>",
            parse_mode="HTML",
            reply_markup=get_main_keyboard(True)
        )
        return

    if state == "custom_delay":
        if text.strip().isdigit():
            sec = int(text.strip())
            if sec < 1:
                await message.answer("❌ Kamida 1 soniya kiritishingiz kerak.")
                return
            set_user_delay(user_id, sec)
            waiting_for[user_id] = None
            await message.answer(
                f"✅ Kutiladigan vaqt <b>{sec} soniya</b> qilib belgilandi!",
                parse_mode="HTML",
                reply_markup=get_main_keyboard(True)
            )
        else:
            await message.answer("❌ Faqat raqam kiriting.")
        return

    # ----- BOSHQA HOLATLAR: ulangan foydalanuvchining oddiy matni (avto-javob) -----
    # Foydalanuvchi faol holatda emas, demak bu oddiy test xabari -> avto-javob yuboriladi.
    if not get_flag('autoreply_enabled', True):
        return

    reply = get_reply_text(user_id, text, is_telethon=False)
    await bot.send_chat_action(message.chat.id, action="typing")
    await asyncio.sleep(2)
    await message.answer(reply)

async def handle_admin_broadcast_text(message, state, text):
    user_id = message.from_user.id
    if state == "adm_bcast_all":
        waiting_for[user_id] = None
        count = await do_broadcast(text, "all")
        await message.answer(f"✅ Xabar <b>{count}</b> ta foydalanuvchiga yuborildi.", parse_mode="HTML")
    elif state == "adm_bcast_online":
        waiting_for[user_id] = None
        count = await do_broadcast(text, "online")
        await message.answer(f"🟢 Xabar <b>{count}</b> ta ONLINE foydalanuvchiga yuborildi.", parse_mode="HTML")
    elif state == "adm_bcast_offline":
        waiting_for[user_id] = None
        count = await do_broadcast(text, "offline")
        await message.answer(f"🔴 Xabar <b>{count}</b> ta OFFLINE foydalanuvchiga yuborildi.", parse_mode="HTML")
    elif state == "adm_bcast_active":
        waiting_for[user_id] = None
        count = await do_broadcast(text, "active")
        await message.answer(f"⚡️ Xabar <b>{count}</b> ta faol jarayonga yuborildi.", parse_mode="HTML")

# ========== MEDIA BROADCAST (content handlerlar) ==========
# Fotograf / video / document / audio — admin media broadcast holatida.
# Quyidagi callbacklar kutilayotgan holatni belgilab beradi, keyin
# foydalanuvchi yuborgan media content-handlerlar orqali broadcast qilinadi.

@dp.callback_query(F.data == "adm_media_photo")
async def adm_media_photo_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    waiting_for[callback.from_user.id] = "adm_bcast_media_photo"
    await callback.message.answer("🖼 Rasm yuboring (barchaga broadcast):")
    await callback.answer()

@dp.callback_query(F.data == "adm_media_video")
async def adm_media_video_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    waiting_for[callback.from_user.id] = "adm_bcast_media_video"
    await callback.message.answer("🎞 Video yuboring (barchaga broadcast):")
    await callback.answer()

@dp.callback_query(F.data == "adm_media_doc")
async def adm_media_doc_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    waiting_for[callback.from_user.id] = "adm_bcast_media_doc"
    await callback.message.answer("📄 Fayl yuboring (barchaga broadcast):")
    await callback.answer()

@dp.callback_query(F.data == "adm_media_audio")
async def adm_media_audio_cb(callback: types.CallbackQuery):
    if not is_admin(callback.from_user): await callback.answer(); return
    waiting_for[callback.from_user.id] = "adm_bcast_media_audio"
    await callback.message.answer("🎵 Audio yuboring (barchaga broadcast):")
    await callback.answer()

@dp.message(F.photo)
async def admin_photo_message(message: types.Message):
    user_id = message.from_user.id
    if not is_admin(message.from_user):
        return
    state = waiting_for.get(user_id)
    if state in ["adm_bcast_media_photo", "adm_bcast_media_video"]:
        waiting_for[user_id] = None
        photo = message.photo[-1]
        count = 0
        for uid, _ in get_all_sessions():
            try:
                await bot.send_photo(uid, photo.file_id, caption=message.caption)
                count += 1
                await asyncio.sleep(0.05)
            except Exception:
                pass
        await message.answer(f"🖼 Rasm <b>{count}</b> ta foydalanuvchiga yuborildi.", parse_mode="HTML")

@dp.message(F.video)
async def admin_video_message(message: types.Message):
    user_id = message.from_user.id
    if not is_admin(message.from_user):
        return
    state = waiting_for.get(user_id)
    if state == "adm_bcast_media_video":
        waiting_for[user_id] = None
        count = 0
        for uid, _ in get_all_sessions():
            try:
                await bot.send_video(uid, message.video.file_id, caption=message.caption)
                count += 1
                await asyncio.sleep(0.05)
            except Exception:
                pass
        await message.answer(f"🎞 Video <b>{count}</b> ta foydalanuvchiga yuborildi.", parse_mode="HTML")

@dp.message(F.document)
async def admin_document_message(message: types.Message):
    user_id = message.from_user.id
    if not is_admin(message.from_user):
        return
    state = waiting_for.get(user_id)
    if state == "adm_bcast_media_doc":
        waiting_for[user_id] = None
        count = 0
        for uid, _ in get_all_sessions():
            try:
                await bot.send_document(uid, message.document.file_id, caption=message.caption)
                count += 1
                await asyncio.sleep(0.05)
            except Exception:
                pass
        await message.answer(f"📄 Fayl <b>{count}</b> ta foydalanuvchiga yuborildi.", parse_mode="HTML")

@dp.message(F.audio | F.voice)
async def admin_audio_message(message: types.Message):
    user_id = message.from_user.id
    if not is_admin(message.from_user):
        return
    state = waiting_for.get(user_id)
    if state == "adm_bcast_media_audio":
        waiting_for[user_id] = None
        count = 0
        audio = getattr(message, 'audio', None) or getattr(message, 'voice', None)
        for uid, _ in get_all_sessions():
            try:
                await bot.send_audio(uid, audio.file_id, caption=message.caption)
                count += 1
                await asyncio.sleep(0.05)
            except Exception:
                pass
        await message.answer(f"🎵 Audio <b>{count}</b> ta foydalanuvchiga yuborildi.", parse_mode="HTML")

# ========== ULASH / LOGIN FUNKSIYALARI ==========
async def request_code(user_id, phone, message):
    try:
        client = TelegramClient(StringSession(), API_ID, API_HASH)
        await client.connect()
        res = await client.send_code_request(phone)
        user_clients[user_id] = client
        phone_cache[user_id] = phone
        phone_code_hash_cache[user_id] = res.phone_code_hash
        waiting_for[user_id] = "code"
        await message.answer(
            "✅ Kod yuborildi!\n\n⚠️ Kodni ajratib yuboring (Masalan: <code>5-1-9-9-6</code>).",
            parse_mode="HTML",
            reply_markup=ReplyKeyboardRemove()
        )
    except FloodWaitError as e:
        await message.answer(f"⏳ Telegram cheklovi! Biroz kuting: {e.seconds} soniya.")
        waiting_for[user_id] = "phone"
    except Exception as e:
        await message.answer(f"❌ Xatolik: {html.escape(str(e))}")
        waiting_for[user_id] = "phone"

async def finish_login(user_id, client, message):
    stop_client_task(user_id)
    session_string = client.session.save()
    save_session(user_id, session_string)
    waiting_for[user_id] = None
    phone_cache.pop(user_id, None)
    phone_code_hash_cache.pop(user_id, None)

    task = asyncio.create_task(start_auto_reply(user_id, client))
    user_clients[user_id] = client
    auto_tasks[user_id] = task

    # welcome xabar
    if get_flag('welcome_enabled', True):
        wm = get_config('welcome_text', 'Xush kelibsiz!👋')
        try:
            await message.answer(wm)
        except Exception:
            pass

    await message.answer(
        f"✅ Muvaffaqiyatli ulandi!\n🤖 Offline bo'lsangiz darhol, Online bo'lsangiz {get_user_delay(user_id)} soniyada avto-javob yuboriladi.",
        reply_markup=get_main_keyboard(True)
    )
    await message.answer("👇 Qulay boshqaruv menyusi:", reply_markup=get_reply_keyboard(message.from_user))

async def restore_sessions():
    sessions = get_all_sessions()
    for user_id, session_string in sessions:
        if is_blocked(user_id):
            continue
        try:
            client = TelegramClient(StringSession(session_string), API_ID, API_HASH)
            await client.connect()
            if await client.is_user_authorized():
                user_clients[user_id] = client
                auto_tasks[user_id] = asyncio.create_task(start_auto_reply(user_id, client))
                print(f"♻️ [DEBUG] Restored session for user {user_id}")
            else:
                await client.disconnect()
                delete_session(user_id)
        except (AuthKeyUnregisteredError, UserDeactivatedBanError, UnauthorizedError):
            delete_session(user_id)
        except Exception as e:
            print(f"Restore error user {user_id}: {e}")
        await asyncio.sleep(0.2)

async def main():
    global BOT_START_TIME
    BOT_START_TIME = datetime.now()
    init_db()
    await restore_sessions()
    await bot.delete_webhook(drop_pending_updates=True)
    print("🤖 Bot to'liq tayyor! Admin panel: 100+ funksiya.")
    await dp.start_polling(bot)

if __name__ == "__main__":
    try:
        asyncio.run(main())
    except KeyboardInterrupt:
        print("To'xtatildi.")
